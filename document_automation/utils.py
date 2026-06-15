import os
import logging
import re
import subprocess

logger = logging.getLogger(__name__)


def process_template(template_path, template_type, input_data):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    if template_type in ("TXT", "HTML"):
        from jinja2 import Environment, FileSystemLoader

        template_dir = os.path.dirname(template_path)
        template_file = os.path.basename(template_path)
        env = Environment(loader=FileSystemLoader(template_dir))
        jinja_template = env.get_template(template_file)

        processed_data = _resolve_clause_refs(input_data or {})
        return jinja_template.render(**processed_data) if processed_data else jinja_template.render()

    elif template_type == "DOCX":
        processed_data = _resolve_clause_refs(input_data or {})

        from docxtpl import DocxTemplate

        doc = DocxTemplate(template_path)
        doc.render(processed_data)
        return doc

    else:
        raise ValueError(f"Unsupported template type: {template_type}")


def _resolve_clause_refs(input_data):
    from document_automation.models import Clause

    resolved = {}
    for key, value in input_data.items():
        if isinstance(value, str) and value.startswith('clause:'):
            clause_id = value.split(':', 1)[1]
            try:
                clause = Clause.objects.get(id=clause_id, is_active=True)
                resolved[key] = clause.content
            except Clause.DoesNotExist:
                resolved[key] = f'[Clause {clause_id} not found]'
        else:
            resolved[key] = value
    return resolved


def generate_output_document(rendered_content, template_type, output_path, output_format='DOCX'):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    if template_type == "DOCX":
        rendered_content.save(output_path)
    elif template_type == "HTML":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(rendered_content)
    elif template_type == "TXT":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(rendered_content)
    else:
        raise ValueError(f"Unsupported template type: {template_type}")

    pdf_output_path = None
    if output_format in ('PDF', 'BOTH'):
        pdf_output_path = convert_to_pdf(output_path)

    return output_path, pdf_output_path


def convert_to_pdf(docx_path):
    pdf_path = docx_path.rsplit('.', 1)[0] + '.pdf'
    try:
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path],
            check=True, capture_output=True, timeout=60
        )
    except (subprocess.CalledProcessError, FileNotFoundError) as e:
        logger.warning(f"PDF conversion failed: {e}")
        return None
    return pdf_path if os.path.exists(pdf_path) else None


def extract_template_fields(template_path, template_type):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    if template_type == "DOCX":
        from docx import Document

        doc = Document(template_path)
        content = '\n'.join(p.text for p in doc.paragraphs)
    else:
        with open(template_path, 'r', encoding='utf-8') as f:
            content = f.read()

    var_pattern = re.findall(r'\{\{\s*(\w+)\s*\}\}', content)
    block_pattern = re.findall(r'\{%\s*(?:for|if|include)\s+(\w+)', content)

    all_vars = set(var_pattern + block_pattern)
    return sorted(all_vars)


def merge_bulk_outputs(output_paths, merged_path):
    from docx import Document
    from copy import deepcopy
    from docx.oxml.ns import qn

    merged_doc = Document()

    for i, path in enumerate(output_paths):
        if i > 0:
            merged_doc.add_page_break()

        doc = Document(path)
        for element in doc.element.body:
            if element.tag != qn('w:sectPr'):
                merged_doc.element.body.append(deepcopy(element))

    for p in merged_doc.paragraphs:
        if p.text == '' and p.runs and not any(attr for attr in dir(p) if 'break' in attr.lower()):
            p._element.getparent().remove(p._element)
            break

    merged_doc.save(merged_path)
    return merged_path


def inject_clause_into_docx(docx_path, clause, insertion_point):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(docx_path)
    body = doc.element.body
    paras = list(body.iterchildren(qn('w:p')))

    if not paras or insertion_point >= len(paras):
        return docx_path

    idx = max(0, min(insertion_point, len(paras) - 1))
    ref_para = paras[idx]

    clause_text = clause.content.strip() if clause.content else ''
    lines = clause_text.split('\n')

    for line in reversed(lines):
        new_para = OxmlElement('w:p')
        new_run = OxmlElement('w:r')
        new_text = OxmlElement('w:t')
        new_text.set(qn('xml:space'), 'preserve')
        new_text.text = line if line else ' '
        new_run.append(new_text)
        new_para.append(new_run)
        ref_para.addnext(new_para)

    doc.save(docx_path)
    return docx_path
