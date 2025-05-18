
import hashlib
import fitz  # PyMuPDF
import json
import re
import openai
import os
from dotenv import load_dotenv
# from langchain_openai import ChatOpenAI

from docx import Document
from django.conf import settings


from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from io import BytesIO

from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document

import tempfile

from md2docx_python.src.md2docx_python import markdown_to_word

from unstructured.partition.md import partition_md
from unstructured.documents.elements import Title, NarrativeText, ListItem

# Load environment variables
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")


# ------------------------------------------
# 📄 Text Extraction from PDF
# ------------------------------------------
def extract_text_from_pdf(pdf_path):
    """Extract text content from each page of a PDF."""
    doc = fitz.open(pdf_path)
    return [page.get_text() for page in doc]

# ------------------------------------------
# 🧠 Generate a unique hash for a page
# ------------------------------------------
def generate_page_hash(text):
    """Generate a simple MD5 hash for a page's text."""
    return hashlib.md5(text.encode('utf-8')).hexdigest()

# ------------------------------------------
# 🧠 Build Prompt: Extract Products & Services
# ------------------------------------------
def build_product_service_prompt(text):
    """Build a prompt for GPT-4o to extract products/services from text."""
    return f"""
You are a business analyst AI. Given the extracted text from a product or service catalogue, identify:
1. Product/service name
2. Detailed description
3. Challenges/problems it solves
4. Target client categories per region: US, UK, South Africa, Europe, Australia

Return ONLY valid JSON structured like this (and NOTHING else):

```json
{{
  "name": "",
  "details": "",
  "solves": "",
  "target_clients": {{
    "US": [],
    "UK": [],
    "South Africa": [],
    "Europe": [],
    "Australia": []
  }}
}}
```

Text:
{text}
"""

# ------------------------------------------
# 🧠 Build Prompt: Product Recommendations based on Persona
# ------------------------------------------
def build_product_recommendation_prompt(persona_json):
    """Builds a prompt for GPT to recommend products/services based on a client persona."""
    return f"""
You are a Sales Enablement AI Assistant.

Based on the following client persona:

{json.dumps(persona_json, indent=2)}

Recommend up to 5 Morae products or services that best fit this client's needs.

For each recommendation, provide:
- Product or Service Name
- Short Reason why it is a good fit

Return the answer ONLY as structured JSON like this:

```json
{{
  "recommendations": [
    {{
      "product_or_service": "Product/Service Name 1",
      "reason": "Why it fits"
    }},
    {{
      "product_or_service": "Product/Service Name 2",
      "reason": "Why it fits"
    }}
  ]
}}
"""


# ------------------------------------------
# 🧠 Build Prompt: Client Persona Generation
# ------------------------------------------
def build_client_persona_prompt(company_name):
    """Build a prompt to generate a client persona based on a company name."""
    return f"""
You are an expert marketing AI. Generate a client persona for the company '{company_name}'.
The persona must include:
- Company Name (Search the internet to see if available. You must not make up a name or generate a fake one)
- Contact Name (Search the internet to see if available. You must not make up a name or generate a fake one)
- Contact Title (Search the internet to see if available. You must not make up a name or generate a fake one)
- Contact Email (Search the internet to see if available. You must not make up a name or generate a fake one)
- Contact Phone (Search the internet to see if available. You must not make up a name or generate a fake one)
- Company Website (Search the internet to see if available. You must not make up a name or generate a fake one)
- Company Description (Search the internet to see if available. You must not make up a name or generate a fake one)
- Industry (Search the internet to see if available. You must not make up a name or generate a fake one)
- Company Size (Search the internet to see if available. You must not make up a name or generate a fake one)
- Key Challenges (Search the internet to see if available. You must not make up a name or generate a fake one)
- Goals (Search the internet to see if available. You must not make up a name or generate a fake one)
- Decision Makers (Search the internet to see if available. You must not make up a name or generate a fake one)
- Region Presence (Search the internet to see if available. You must not make up a name or generate a fake one)
- Technologies Used (Search the internet to see if available. You must not make up a name or generate a fake one)
- Buying Triggers (Search the internet to see if available. You must not make up a name or generate a fake one)

Return the output as **standardized JSON** only:

```json
{{
  "company_name": "",
  "contact_name": "",
  "contact_title": "",
  "contact_email": "",
  "contact_phone": "",
  "company_website": "",
  "company_description": "",
  "industry": "",
  "company_size": "",
  "key_challenges": [],
  "goals": [],
  "decision_makers": [],
  "regions": [],
  "technologies_used": [],
  "buying_triggers": []
}}
```
"""

# ------------------------------------------
# 🧠 Build Prompt: Suggest Products for Persona
# ------------------------------------------
def build_suggested_solutions_prompt(persona_json, catalog_json):
    return f"""
You are a B2B Sales AI.

Here is a prospective client persona:
{json.dumps(persona_json, indent=2)}

Here is Morae's full catalog of services:
{json.dumps(catalog_json, indent=2)}

📋 Your task:
- Analyze the persona's needs and match them to Morae's solutions.
- Recommend up to **5** Morae products or services most relevant to the client.
- For **each match**, include:
  - Product/Service Name
  - Reason why it fits the client's challenges or goals.

🧠 Output Format:
Return **ONLY strict JSON**, with no explanations, no markdown, no commentary.
Structure exactly like this:

```json
{{
  "recommendations": [
    {{
      "product_or_service": "Product Name 1",
      "reason": "Why this product is a good fit."
    }},
    {{
      "product_or_service": "Product Name 2",
      "reason": "Why this product is a good fit."
    }}
  ]
}}
```
"""

# ------------------------------------------
# 🧠 Build Prompt: Draft Proposal
# ------------------------------------------
def build_proposal_prompt(persona_json, additional_info="", focus_solution=None, context_solutions=None):
    """
    Builds a detailed proposal prompt for GPT, optionally focused on a specific product or service.
    """
    prompt = f"""
You are a Morae Global staff and expert legal contract AI specialized in drafting professional business documents.

Generate a highly professional solution(s) proposal from Morae Global for the following client persona:

{json.dumps(persona_json, indent=2)}

"""

    if focus_solution:
        prompt += f"""
🎯 Focus Solution:
This proposal should focus primarily on the following product or service:

- **Name:** {focus_solution.get("product_or_service", "")}
- **Why it fits:** {focus_solution.get("reason_for_relevance", "")}
"""

    if context_solutions:
        prompt += f"""

📚 Context Solutions:
These additional offerings may complement or provide alternative value:

{json.dumps(context_solutions, indent=2)}
"""

    prompt += f"""

📝 Additional Notes:
{additional_info}

Formatting Instructions:
- Use proper Heading hierarchy: 
  - H1 for major sections (e.g., Introduction, Objectives)
  - H2 for subsections (e.g., Detailed Goals)
  - H3 for sub-subsections if needed
- Use hierarchical numbering for sections and subsections:
  - 1. Main Section
    - 1.1 Subsection
    - 1.1.1 Sub-subsection
- Use bullet points `-` for unordered lists where appropriate
- Use Bold for key titles, section names, or highlights
- Use Italics for emphasis or side notes
- Maintain professional tone, business English
- Add logical paragraph breaks after headings or bullets
- NO Markdown artifacts like '---' or triple backticks (` ``` `)
- Write fully structured readable proposal text, NOT JSON or code blocks.

Start with a title page section.

Ensure the document is structured cleanly for easy DOCX export later.
"""

    return prompt


# ------------------------------------------
# 🚀 GPT-4o Call (Upgraded)
# ------------------------------------------
def call_gpt4o(prompt):
    """Call GPT-4o API safely and reliably to complete the prompt and extract JSON or text."""
    
    from openai import OpenAI
    client = OpenAI()

    try:
        response = client.responses.create(
            model="gpt-4o",
            tools=[{"type": "web_search_preview"}],
            input=prompt
        )
        # content = response.choices[0].message.content.strip()
        
        content = response.output_text

        # Try extracting JSON if possible
        json_match = re.search(r'\{.*\}', content, re.DOTALL)

        if json_match:
            json_string = json_match.group(0)

            try:
                return json.loads(json_string)
            except json.JSONDecodeError as json_err:
                print(f"⚠️ JSON decoding failed, returning text instead: {json_err}")
                return content

        # If no JSON found, return content as plain text
        return content

    except Exception as e:
        print(f"⚠️ GPT-4o call failed: {e}")
        return {}
    

def save_proposal_to_docx(proposal):
    """Save the ProposalDraft content into a DOCX file nicely formatted with heading hierarchy, logo, footer."""
    persona = proposal.persona
    persona_data = persona.standardized_persona

    # Create a new Word Document
    doc = Document()

    # Add Logo
    logo_path = os.path.join(settings.MEDIA_ROOT, "branding", "logo.png")  # customize your logo path
    if os.path.exists(logo_path):
        header = doc.sections[0].header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(2))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Title
    doc.add_heading('Proposal', 0)

    # Client Persona Overview
    doc.add_heading('Client Persona Overview', level=1)
    doc.add_paragraph(f"Company Name: {persona_data.get('company_name', '')}")
    doc.add_paragraph(f"Contact Name: {persona_data.get('contact_name', '')}")
    doc.add_paragraph(f"Contact Title: {persona_data.get('contact_title', '')}")
    doc.add_paragraph(f"Contact Email: {persona_data.get('contact_email', '')}")
    doc.add_paragraph(f"Contact Phone: {persona_data.get('contact_phone', '')}")
    doc.add_paragraph(f"Company Website: {persona_data.get('company_website', '')}")
    doc.add_paragraph(f"Company Description: {persona_data.get('company_description', '')}")
    doc.add_paragraph(f"Industry: {persona_data.get('industry', '')}")
    doc.add_paragraph(f"Company Size: {persona_data.get('company_size', '')}")
    doc.add_paragraph(f"Regions: {', '.join(persona_data.get('regions', []))}")
    doc.add_paragraph(f"Key Challenges: {', '.join(persona_data.get('key_challenges', []))}")
    doc.add_paragraph(f"Goals: {', '.join(persona_data.get('goals', []))}")
    doc.add_paragraph(f"Buying Triggers: {', '.join(persona_data.get('buying_triggers', []) )}")
    doc.add_paragraph(f"Decision Makers: {', '.join(persona_data.get('decision_makers', []))}")
    doc.add_paragraph(f"Technologies Used: {', '.join(persona_data.get('technologies_used', []))}")

    # Spacer
    doc.add_page_break()

    # Proposal Content Parsing
    lines = proposal.content.splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")  # Empty line
        elif re.match(r'^\d+\.$', line):  # Top-level numbered heading like 1.
            doc.add_heading(line, level=1)
        elif re.match(r'^\d+\.\d+\.$', line):  # Second level like 1.1.
            doc.add_heading(line, level=2)
        elif re.match(r'^\d+\.\d+\.\d+\.$', line):  # Third level like 1.1.1.
            doc.add_heading(line, level=3)
        elif line.startswith("- ") or line.startswith("* "):  # Bullet list
            para = doc.add_paragraph(style='ListBullet')
            para.add_run(line[2:].strip())
        elif re.match(r'^\d+\) ', line):  # Numbered list like 1) 2) 3)
            para = doc.add_paragraph(style='ListNumber')
            para.add_run(line.split(' ', 1)[1])
        else:
            # Normal paragraph with minor inline formatting
            para = doc.add_paragraph()
            run = para.add_run()

            # Bold
            if "**" in line:
                line = re.sub(r'\*\*(.*?)\*\*', lambda m: f"{m.group(1).upper()}", line)

            # Italics
            if "_" in line:
                line = re.sub(r'_(.*?)_', lambda m: f"{m.group(1)}", line)

            run.add_text(line)
            run.font.size = Pt(11)

    # Add Footer
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Generated by Morae Sales Enablement Platform"
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save to media/proposal_docs/
    proposal_dir = os.path.join(settings.MEDIA_ROOT, "proposal_docs")
    os.makedirs(proposal_dir, exist_ok=True)

    filepath = os.path.join(proposal_dir, f"proposal_{proposal.id}.docx")
    doc.save(filepath)

    # Save the file path to the ProposalDraft model
    proposal.docx_file.name = os.path.relpath(filepath, settings.MEDIA_ROOT)
    proposal.save()


def build_solution_chat_prompt(solution, question, context=None):
    context_section = f"\nPrevious chat context:\n{context}" if context else ""

    return f"""
You are a Sales Assistant AI.

Here is a product or service:
- Name: {solution.product_or_service}
- Why it is recommended: {solution.reason_for_relevance}

{context_section}

A sales team member has this question:
"{question}"

Answer helpfully, with context on why this solution is a good fit for the given persona.
"""


def generate_proposal_docx_from_markdown(company_name: str, notes: str, markdown_content: str) -> BytesIO:
    # Construct the full markdown content
    full_markdown = f"""
# Proposal

**Issued by:** Morae Global  
**For:** {company_name}  
**Date:** [Insert Date]

---

## 1. Introduction

{markdown_content}

---

## Notes

{notes}
    """

    # Write markdown to a temporary .md file
    with tempfile.NamedTemporaryFile(mode='w+', delete=False, suffix='.md') as md_file:
        md_file.write(full_markdown)
        md_file_path = md_file.name

    # Prepare output .docx file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as docx_file:
        docx_path = docx_file.name

    # Convert markdown to Word
    markdown_to_word(md_file_path, docx_path)

    # Load the .docx into a BytesIO buffer
    with open(docx_path, 'rb') as f:
        buffer = BytesIO(f.read())

    buffer.seek(0)
    return buffer


# ------------------------------------------
# 🧠 GPT Prompt: Extract RFP Sections
# ------------------------------------------
def build_rfp_extraction_prompt(text):
    """
    Builds a GPT prompt to extract structured sections from an uploaded client RFP.
    """
    return f"""
You are a legal AI assistant analyzing a client's Request for Proposal (RFP).

From the RFP text below, extract the following key sections:
- Executive Summary (1-2 paragraphs)
- Client Requirements
- Timeline
- Evaluation Criteria
- Submission Instructions
- Contact Information
- Any Explicit Questions (list them)

🧠 Output Format:
Return JSON only — no markdown, no commentary.

```json
{{
  "summary": "...",
  "parsed_sections": {{
    "requirements": "...",
    "timeline": "...",
    "evaluation_criteria": "...",
    "submission_instructions": "...",
    "contact_info": "...",
    "questions": [
      "Question 1?",
      "Question 2?"
    ]
  }}
}}
```
Text:
{text}

"""

# ------------------------------------------
# 🧠 Extract RFP Sections
# ------------------------------------------
def extract_sections_from_rfp(text):
    """
    Extract structured sections from the RFP text using GPT-4o.
    """
    prompt = build_rfp_extraction_prompt(text)
    response = call_gpt4o(prompt)

    if isinstance(response, dict) and "parsed_sections" in response:
        return response["parsed_sections"]
    else:
        print("⚠️ Failed to extract sections from RFP.")
        return {}
    

# ------------------------------------------
# 🧠 GPT Prompt: Answer RFP Chat Questions
# ------------------------------------------
def build_rfp_chat_prompt(rfp_summary, question, history=None):
    """
    Builds a GPT prompt for answering a question about an RFP.
    """
    context = f"\n\nPrevious Chat Context:\n{history}" if history else ""

    return f"""
You are a sales assistant AI helping respond to an RFP.

📄 Here is a summary and key sections extracted from the RFP:

{json.dumps(rfp_summary, indent=2)}

💬 A sales team member asked:
{question}

{context}

Answer clearly, referencing relevant RFP sections and keeping the response brief but helpful.
"""

# ------------------------------------------
# 🧠 Try to parse JSON from response content
# ------------------------------------------
def try_parse_json_from_response(content):
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        return None


# ------------------------------------------
# 🧠 GPT Prompt: Extract RFP Metadata (Client Info, Due Date etc.)
# ------------------------------------------
def build_rfp_metadata_prompt(text):
    return f"""
You are a legal AI assistant helping a sales team process incoming RFPs.

From the following text, extract this information in strict JSON format:
- Client Name
- RFP Title (if available)
- Submission Due Date
- Country or Region
- Contact Person and Email (if available)

Return strictly JSON only:

```json
{{
  "client_name": "",
  "rfp_title": "",
  "due_date": "",
  "region": "",
  "contact_name": "",
  "contact_email": ""
}}
```
Text:
{text}
"""

# ------------------------------------------
# 🧠 Extract RFP Metadata
# ------------------------------------------
def extract_rfp_metadata(text: str) -> dict:
    prompt = build_rfp_metadata_prompt(text)
    response = call_gpt4o(prompt)
        
    if isinstance(response, dict):
        return response
    else:
        parsed = try_parse_json_from_response(response)
        return parsed if parsed else {}
    

