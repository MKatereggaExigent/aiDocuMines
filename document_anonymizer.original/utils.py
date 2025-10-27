import os
import json
import re
import logging
from django.conf import settings
from docx import Document
import spacy
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from unstructured.partition.pdf import partition_pdf
from document_anonymizer.models import Anonymize

from django.db.models import Avg, Count, Q
from document_anonymizer.models import AnonymizationRun, Anonymize, DeAnonymize

from functools import lru_cache

logger = logging.getLogger(__name__)


class AnonymizationService:
    """
    Handles document extraction, full Presidio ➔ SpaCy anonymization pipeline, and de-anonymization.
    Now supports structured PDF extraction using unstructured.
    """

    def __init__(self):
        logger.info("📦 Initializing AnonymizationService with shared models...")
        self.analyzer = AnalyzerEngine()
        self.anonymizer = AnonymizerEngine()
        self.nlp = spacy.load("en_core_web_lg")  # Load once per task

    def extract_text_from_file(self, file_path):
        logger.info(f"🔄 Extracting text from: {file_path}")

        if not os.path.exists(file_path):
            logger.error(f"❌ File not found: {file_path}")
            return None, [], []

        try:
            if file_path.lower().endswith(".pdf"):
                return self.extract_text_from_pdf(file_path)
            elif file_path.lower().endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read(), [], []
            elif file_path.lower().endswith(".docx"):
                return self.extract_text_from_docx(file_path), [], []
            else:
                logger.warning(f"⚠️ Unsupported file type: {file_path}")
                return None, [], []
        except Exception as e:
            logger.error(f"❌ Failed to extract text from {file_path}: {e}")
            return None, [], []

    def extract_structured_text(self, file_path):
        """
        Dynamically select the correct structured extraction function based on file type.
        """
        logger.info(f"📄 extract_structured_text() ➔ Processing {file_path}")
        try:
            if file_path.lower().endswith(".pdf"):
                from unstructured.partition.pdf import partition_pdf
                elements = partition_pdf(filename=file_path)
            elif file_path.lower().endswith(".docx"):
                from unstructured.partition.docx import partition_docx
                elements = partition_docx(filename=file_path)
            else:
                logger.warning(f"⚠️ Unsupported file type for structured extraction: {file_path}")
                return None, [], []

            structured_text = "\n\n".join(el.text.strip() for el in elements if el.text and el.text.strip())
            return structured_text, elements, [el.to_dict() for el in elements]
        except Exception as e:
            logger.error(f"❌ extract_structured_text() failed: {e}")
            return None, [], []


    def extract_text_from_pdf(self, file_path):
        try:
            logger.info(f"📄 Using unstructured to extract structured PDF text from: {file_path}")
            elements = partition_pdf(filename=file_path)
            structured_text = "\n\n".join(el.text.strip() for el in elements if el.text and el.text.strip())

            return structured_text, elements, [el.to_dict() for el in elements]
        except Exception as e:
            logger.error(f"❌ Structured PDF extraction failed via unstructured: {e}")
            return None, [], []

    def extract_text_from_docx(self, file_path):
        try:
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"❌ DOCX text extraction failed: {e}")
            return None
    
    def anonymize_text(self, text):
        """
        Anonymizes the input text using both Presidio and SpaCy engines.
        Reuses shared engine instances for performance and memory efficiency.
        """
        logger.info("🔄 Running Presidio anonymization...")
        presidio_results = self.analyzer.analyze(text=text, entities=[], language="en")
        presidio_map = {}
        presidio_counter = {}
        presidio_masked = text
    
        for result in presidio_results:
            entity_type = result.entity_type
            presidio_counter[entity_type] = presidio_counter.get(entity_type, 0) + 1
            mask = f"{entity_type}_MASKED_{presidio_counter[entity_type]}"
            value = text[result.start:result.end].strip()
            presidio_map[mask] = value
            presidio_masked = re.sub(r'\b' + re.escape(value) + r'\b', mask, presidio_masked, 1)
    
        logger.info("🔄 Running SpaCy anonymization...")
        spacy_map = {}
        spacy_counter = {}
        final_masked = presidio_masked
        doc = self.nlp(presidio_masked)
    
        for ent in doc.ents:
            entity_type = ent.label_
            spacy_counter[entity_type] = spacy_counter.get(entity_type, 0) + 1
            mask = f"{entity_type}_MASKED_{spacy_counter[entity_type]}"
            spacy_map[mask] = ent.text.strip()
            final_masked = re.sub(r'\b' + re.escape(ent.text.strip()) + r'\b', mask, final_masked, 1)
    
        combined_map = {**spacy_map, **presidio_map}
    
        logger.info(f"✅ Presidio map: {json.dumps(presidio_map, indent=2)}")
        logger.info(f"✅ SpaCy map: {json.dumps(spacy_map, indent=2)}")
    
        return final_masked, combined_map, presidio_map, spacy_map


    def deanonymize_pipeline(self, masked_text, spacy_mapping, presidio_mapping):
        partially_restored = self.reverse_masking(masked_text, spacy_mapping)
        fully_restored = self.reverse_masking(partially_restored, presidio_mapping)
        return fully_restored


    def reverse_masking(self, text, entity_mapping):
        for mask, original in sorted(entity_mapping.items(), key=lambda x: len(x[0]), reverse=True):
            pattern = re.compile(r'\b' + re.escape(mask) + r'\b')
            text = pattern.sub(original, text)
        return text


    def get_file_path(self, original_file_path, folder="anonymized", extension="txt"):
        directory = os.path.join(os.path.dirname(original_file_path), folder)
        os.makedirs(directory, exist_ok=True)
        filename = os.path.basename(original_file_path).replace(".pdf", f".{extension}")
        return os.path.join(directory, filename)

    def update_block_text(self, structured_blocks, element_id, new_text):
        for block in structured_blocks:
            if block.get("element_id") == element_id:
                block["text"] = new_text
                return True
        return False

    '''
    def compute_risk_score(self, original_text: str, presidio_map: dict, spacy_map: dict):
        """
        Comprehensive PII risk scoring using Presidio + spaCy.

        - De-dupes by (normalized_value, label) across engines
        - Weights labels by severity (gov IDs/financial IDs > medical > contact > generic)
        - Scales by density per 1,000 tokens (document-size proportional)
        - Covers full spaCy NER (OntoNotes) + the Presidio entities you enumerated
        - Still resilient to unknown/custom labels (sane defaults)

        Returns:
          {
            "risk_score": float (0..100),
            "risk_level": "Ok" | "Low" | "Medium" | "High",
            "breakdown": { label: count }
          }
        """
        import re

        text = (original_text or "").strip()
        # token count (fallback to 1 to avoid div-by-zero)
        doc_tokens = max(1, len(re.findall(r"\w+", text)))

        # -------- helpers --------
        def label_of(mask_key: str) -> str:
            # e.g. "EMAIL_ADDRESS_MASKED_3" -> "EMAIL_ADDRESS"
            return (mask_key or "").split("_MASKED_")[0].upper().strip()

        def norm_val(v: str) -> str:
            # normalize entity surface form for de-duping
            v = (v or "").strip().lower()
            v = re.sub(r"\s+", " ", v)
            return v

        # -------- spaCy labels (canonical + common alternates) --------
        SPACY_STD = {
            "PERSON", "NORP", "FAC", "ORG", "GPE", "LOC", "PRODUCT", "EVENT",
            "WORK_OF_ART", "LAW", "LANGUAGE", "DATE", "TIME", "PERCENT", "MONEY",
            "QUANTITY", "ORDINAL", "CARDINAL"
        }
        # multilingual / alternate label forms you mentioned
        SPACY_ALIASES = {
            "PER": "PERSON",
            "ORG": "ORG",
            "LOC": "LOC",
            "MISC": "MISC",
        }

        # -------- Presidio entities (from your lists) --------
        # Finance / credentials
        PRES_FINANCE_STRICT = {
            "CREDIT_CARD", "US_BANK_NUMBER", "IBAN_CODE", "CRYPTO",
        }
        # Contact & network
        PRES_CONTACT = {
            "EMAIL_ADDRESS", "PHONE_NUMBER", "IP_ADDRESS", "URL",
        }
        # Generic person/location (Presidio’s logical entities)
        PRES_GENERIC = {"PERSON", "LOCATION", "NRP"}  # NRP ~ NORP
        # Medical
        PRES_MEDICAL = {"MEDICAL_LICENSE", "UK_NHS", "AU_MEDICARE"}
        # US IDs
        PRES_US = {"US_SSN", "US_ITIN", "US_PASSPORT", "US_DRIVER_LICENSE"}
        # UK
        PRES_UK = {"UK_NHS", "UK_NINO"}
        # Spain
        PRES_ES = {"ES_NIF", "ES_NIE"}
        # Italy
        PRES_IT = {"IT_FISCAL_CODE", "IT_DRIVER_LICENSE", "IT_VAT_CODE", "IT_PASSPORT", "IT_IDENTITY_CARD"}
        # Poland
        PRES_PL = {"PL_PESEL"}
        # Singapore
        PRES_SG = {"SG_NRIC_FIN", "SG_UEN"}
        # Australia
        PRES_AU = {"AU_ABN", "AU_ACN", "AU_TFN", "AU_MEDICARE"}
        # India
        PRES_IN = {"IN_PAN", "IN_AADHAAR", "IN_VEHICLE_REGISTRATION", "IN_VOTER", "IN_PASSPORT"}
        # Finland
        PRES_FI = {"FI_PERSONAL_IDENTITY_CODE"}
        # Korea
        PRES_KR = {"KR_RRN"}
        # Thailand
        PRES_TH = {"TH_TNIN"}
        # Date/time (Presidio combined)
        PRES_DATETIME = {"DATE_TIME"}

        # -------- severity groups (weights) --------
        # High severity: government/strong identifiers and financial credentials
        HIGH = (
            PRES_FINANCE_STRICT
            | PRES_US | PRES_UK | PRES_ES | PRES_IT | PRES_PL | PRES_SG
            | PRES_AU | PRES_IN | PRES_FI | PRES_KR | PRES_TH
            | {"IT_PASSPORT", "IN_PASSPORT"}  # already included above; kept for clarity
        )
        # Medium-High: medical identifiers (very sensitive); some org ID numbers
        MED_HIGH = PRES_MEDICAL | {"SG_UEN", "AU_ABN", "AU_ACN", "IT_VAT_CODE"}
        # Medium: direct contact/trackers + demographic groupings
        MED = PRES_CONTACT | {"NRP", "NORP"}
        # Low+: generic entities and temporal/values
        LOW = (
            PRES_GENERIC
            | PRES_DATETIME
            | {
                "PERSON", "GPE", "LOC", "ORG", "FAC", "PRODUCT", "EVENT", "WORK_OF_ART",
                "LAW", "LANGUAGE", "DATE", "TIME", "PERCENT", "MONEY",
                "QUANTITY", "ORDINAL", "CARDINAL", "MISC"
            }
        )

        def canonicalize_label(raw: str) -> str:
            L = (raw or "").upper().strip()
            # unify spaCy alternates (PER→PERSON etc.)
            if L in SPACY_ALIASES:
                return SPACY_ALIASES[L]
            return L

        def weight_for(label: str) -> float:
            L = canonicalize_label(label)
            if L in HIGH:      return 8.0
            if L in MED_HIGH:  return 6.0
            if L in MED:       return 4.5
            if L in LOW:       return 2.5
            # Fallback heuristics (catch unknown/custom labels)
            if re.search(r"(SSN|PASSPORT|DRIVER|LICENSE|NINO|NRIC|PESEL|TFN|AADHAAR|RRN|TNIN|ID(ENTITY)?|IBAN|BANK|ACCOUNT|CREDIT|CARD|CRYPTO)", L):
                return 7.0
            if re.search(r"(EMAIL|PHONE|MOBILE|TEL|IP|URL|DOMAIN)", L):
                return 4.0
            if re.search(r"(DATE|TIME|MONEY|PERCENT|QUANTITY|ORDINAL|CARDINAL)", L):
                return 2.0
            return 3.0  # sensible default

        # -------- merge & de-duplicate across engines --------
        merged = {}  # (value_norm, label) -> entry
        def ingest(src_map: dict):
            for mask, val in (src_map or {}).items():
                raw_label = label_of(mask)
                label = canonicalize_label(raw_label)
                v_norm = norm_val(val)
                if not v_norm:
                    continue
                key = (v_norm, label)
                entry = merged.setdefault(key, {
                    "value": v_norm,
                    "label": label,
                    "count": 0,
                    "weight": weight_for(label),
                    "tokens": len(re.findall(r"\w+", v_norm)),
                    "chars": len(v_norm),
                })
                entry["count"] += 1

        ingest(presidio_map or {})
        ingest(spacy_map or {})

        # -------- aggregate & scale --------
        total_occurrences = sum(e["count"] for e in merged.values())
        total_weighted    = sum(e["weight"] * e["count"] for e in merged.values())

        # density per 1k tokens => doc-size proportional
        density_per_1k  = (total_occurrences / doc_tokens) * 1000.0
        severity_per_1k = (total_weighted    / doc_tokens) * 1000.0

        # Blend density (60%) + severity (40%), each with a soft cap
        risk = (
            min(1.0, density_per_1k  / 20.0) * 0.60 +
            min(1.0, severity_per_1k / 30.0) * 0.40
        ) * 100.0
        risk_score = round(min(100.0, risk), 2)

        if   risk_score >= 60: risk_level = "High"
        elif risk_score >= 20: risk_level = "Medium"
        elif risk_score >  0:  risk_level = "Low"
        else:                  risk_level = "Ok"

        breakdown = {}
        for e in merged.values():
            breakdown[e["label"]] = breakdown.get(e["label"], 0) + e["count"]

        return {
            "risk_score": risk_score,
            "risk_level": risk_level,
            "breakdown": dict(sorted(breakdown.items(), key=lambda kv: (-kv[1], kv[0]))),
        }
    '''


    def compute_risk_score(self, original_text: str, presidio_map: dict, spacy_map: dict):
        """
        Mask-aware, size-fair PII risk scoring.

        Key ideas:
          - Ignore spaCy hits whose *value* contains Presidio masks (e.g. "..._MASKED_7")
            to avoid label pollution like ORG=“US_DRIVER_LICENSE_MASKED_2...”.
          - De-dupe by (normalized value, label) across engines.
          - Count per-label with diminishing returns (1st hit counts most, repeats count less).
          - Normalize by document length (per ~1k tokens) with a floor so short docs aren’t over-penalized.
          - Map to 0..100 via a logistic curve, so typical docs don’t hit 100.

        Returns:
          {
            "risk_score": float,
            "risk_level": "Ok" | "Low" | "Medium" | "High",
            "breakdown": { label: count }
          }
        """
        import re
        import math

        text = (original_text or "").strip()
        # token count (fallback 1 to avoid div-by-zero)
        doc_tokens = max(1, len(re.findall(r"\w+", text)))
        # size normalization (per ~1k tokens) but don't penalize tiny docs too much
        tokens_k = max(0.5, doc_tokens / 1000.0)

        # ---------- helpers -------------------------------------------------------
        def label_of(mask_key: str) -> str:
            return (mask_key or "").split("_MASKED_")[0].upper().strip()

        def norm_val(v: str) -> str:
            v = (v or "").strip()
            v = re.sub(r"\s+", " ", v)
            return v

        # unify common spaCy aliases
        SPACY_ALIASES = {"PER": "PERSON", "MISC": "MISC", "LOC": "LOC", "ORG": "ORG"}
        def canon_spacy_label(lbl: str) -> str:
            L = (lbl or "").upper().strip()
            return SPACY_ALIASES.get(L, L)

        # ---------- severity taxonomy (sync with risk_weight_for_label) -----------
        HIGH = {
            "CREDIT_CARD","US_BANK_NUMBER","IBAN_CODE","CRYPTO",
            "US_SSN","US_ITIN","US_PASSPORT","US_DRIVER_LICENSE",
            "UK_NINO","UK_NHS",
            "ES_NIF","ES_NIE",
            "IT_FISCAL_CODE","IT_DRIVER_LICENSE","IT_VAT_CODE","IT_PASSPORT","IT_IDENTITY_CARD",
            "PL_PESEL",
            "SG_NRIC_FIN","SG_UEN",
            "AU_ABN","AU_ACN","AU_TFN","AU_MEDICARE",
            "IN_PAN","IN_AADHAAR","IN_VEHICLE_REGISTRATION","IN_VOTER","IN_PASSPORT",
            "FI_PERSONAL_IDENTITY_CODE","KR_RRN","TH_TNIN",
        }
        MED_HIGH = {"MEDICAL_LICENSE","UK_NHS","AU_MEDICARE","SG_UEN","AU_ABN","AU_ACN","IT_VAT_CODE"}
        MED = {"EMAIL_ADDRESS","PHONE_NUMBER","IP_ADDRESS","URL","NRP","NORP"}
        LOW = {
            "PERSON","GPE","LOC","ORG","FAC","PRODUCT","EVENT","WORK_OF_ART","LAW","LANGUAGE",
            "DATE","TIME","PERCENT","MONEY","QUANTITY","ORDINAL","CARDINAL","MISC","LOCATION","DATE_TIME"
        }

        def weight_for(label: str) -> float:
            L = label.upper().strip()
            if L in HIGH: return 8.0
            if L in MED_HIGH: return 6.0
            if L in MED: return 4.5
            if L in LOW: return 2.5
            # Fallback heuristics for unknown/custom labels
            if re.search(r"(SSN|PASSPORT|DRIVER|LICENSE|NINO|NRIC|PESEL|TFN|AADHAAR|RRN|TNIN|IBAN|BANK|ACCOUNT|CREDIT|CARD|CRYPTO)", L):
                return 7.0
            if re.search(r"(EMAIL|PHONE|MOBILE|TEL|IP|URL|DOMAIN)", L):
                return 4.0
            if re.search(r"(DATE|TIME|MONEY|PERCENT|QUANTITY|ORDINAL|CARDINAL)", L):
                return 2.0
            return 3.0

        def group_k_for(label: str) -> float:
            """
            Diminishing-returns 'k' per severity group.
            Smaller k => saturates faster (first hit counts most).
            """
            w = weight_for(label)
            if w >= 8.0:  # HIGH
                return 1.0
            if w >= 6.0:  # MED_HIGH
                return 1.5
            if w >= 4.5:  # MED
                return 2.0
            return 3.0     # LOW/other

        # ---------- mask-aware ingestion -----------------------------------------
        MASK_TOKEN_RE = re.compile(r"\b([A-Z][A-Z_]+)_MASKED_\d+\b")

        merged = {}  # key=(value_norm, label) -> entry

        def add_entry(v_norm: str, label: str):
            key = (v_norm, label)
            e = merged.setdefault(key, {
                "value": v_norm,
                "label": label,
                "count": 0,
                "weight": weight_for(label),
            })
            e["count"] += 1

        # 1) Presidio: always count original values with their own labels
        for mask, val in (presidio_map or {}).items():
            L = label_of(mask)
            v_norm = norm_val(val)
            if not v_norm:
                continue
            add_entry(v_norm, L)

        # 2) spaCy: ignore any hit whose 'value' contains Presidio mask tokens
        for mask, val in (spacy_map or {}).items():
            raw_v = norm_val(val)
            if not raw_v:
                continue
            # If spaCy value includes any *_MASKED_* tokens, it's derived from Presidio replacement -> skip
            if "_MASKED_" in raw_v:
                continue
            L = canon_spacy_label(label_of(mask))
            add_entry(raw_v, L)

        # ---------- aggregate per label with diminishing returns ------------------
        label_counts = {}
        for e in merged.values():
            label_counts[e["label"]] = label_counts.get(e["label"], 0) + e["count"]

        # Diminishing-returns curve per label: f(n) = 1 - exp(-n/k)
        # Then normalize by tokens_k (≈ per 1k tokens) and weight by severity.
        raw = 0.0
        for label, n in label_counts.items():
            k = group_k_for(label)
            diminishing = 1.0 - math.exp(-float(n) / k)
            raw += (weight_for(label) * diminishing) / tokens_k

        # ---------- logistic map to 0..100 (calibrated) --------------------------
        # Tunables (alpha: steepness, beta: mid-point). Adjust if you want stricter/looser scoring.
        alpha = 0.85
        beta = 4.0
        risk_score = 100.0 / (1.0 + math.exp(-alpha * (raw - beta)))
        risk_score = round(risk_score, 2)

        if   risk_score >= 70.0: risk_level = "High"
        elif risk_score >= 35.0: risk_level = "Medium"
        elif risk_score >  0.0:  risk_level = "Low"
        else:                    risk_level = "Ok"

        # breakdown = actual counted occurrences per canonical label
        breakdown = dict(sorted(label_counts.items(), key=lambda kv: (-kv[1], kv[0])))

        return {
            "risk_score": risk_score,
            "risk_level": risk_level,
            "breakdown": breakdown,
            # (Optional) diagnostics you can log if needed:
            # "doc_tokens": doc_tokens, "tokens_k": tokens_k, "raw": round(raw, 3)
        }


def generate_anonymized_html(masked_text, entity_mapping):
    html_template = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
    <title>Anonymized Document</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            margin: 40px;
            background-color: #f7f9fc;
            color: #2e2e2e;
            line-height: 1.7;
        }
        h1, h2, h3 {
            color: #1a1a1a;
            margin-top: 2rem;
        }
        .masked {
            background-color: #fff3cd;
            color: #856404;
            font-weight: bold;
            padding: 2px 6px;
            border-radius: 4px;
            border: 1px solid #ffeeba;
            cursor: help;
        }
        pre {
            background-color: #ffffff;
            padding: 20px;
            border: 1px solid #ddd;
            border-radius: 6px;
            overflow-x: auto;
            white-space: pre-wrap;
            word-break: break-word;
        }
        .container {
            max-width: 960px;
            margin: auto;
        }
        .toggle-map {
            margin-top: 10px;
            cursor: pointer;
            color: #007bff;
            text-decoration: underline;
        }
        #entity-map {
            display: none;
            margin-top: 10px;
        }
        footer {
            margin-top: 60px;
            font-size: 0.9rem;
            text-align: center;
            color: #888;
        }
    </style>
    <script>
        function toggleEntityMap() {
            var map = document.getElementById('entity-map');
            map.style.display = map.style.display === 'none' ? 'block' : 'none';
        }
    </script>
</head>
<body>
    <div class="container">
        <h1>🛡️ Anonymized Document</h1>
        <pre>{text}</pre>
        <h3 class="toggle-map" onclick="toggleEntityMap()">📜 Show/Hide Entity Mapping</h3>
        <pre id="entity-map">{entity_map}</pre>
    </div>
    <footer>
        Generated by aiDocuMines &mdash; Preserving privacy, securely and beautifully.
    </footer>
</body>
</html>
"""
    highlighted_text = masked_text
    for mask, original in sorted(entity_mapping.items(), key=lambda x: len(x[0]), reverse=True):
        tooltip = f"title='Original: {original}'"
        highlighted_text = highlighted_text.replace(
            mask, f'<span class="masked" {tooltip}>{mask}</span>'
        )

    entity_map_json = json.dumps(entity_mapping, indent=4)
    return html_template.replace("{text}", highlighted_text).replace("{entity_map}", entity_map_json)

from docx import Document as DocxDocument

def summarize_blocks(blocks, element_ids, model="gpt-4"):
    """
    Dummy implementation for now. Replace with actual GPT-4 call.
    """
    results = []
    for block in blocks:
        if block.get("element_id") in element_ids:
            text = block.get("text", "")
            results.append({
                "element_id": block.get("element_id"),
                "summary": f"Summary of: {text[:60]}..."
            })
    return results


def update_structured_block(json_path, element_id, new_text):
    with open(json_path, "r", encoding="utf-8") as f:
        blocks = json.load(f)

    updated = False
    for block in blocks:
        if block.get("element_id") == element_id:
            block["text"] = new_text
            updated = True
            break

    if updated:
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(blocks, f, indent=2)

    return updated


def export_to_markdown(json_path, generate_preview=False):
    """
    Converts structured JSON blocks into a well-formatted Markdown document.
    - Adds a TOC (Table of Contents) based on Titles and Headers.
    - Formats tables and content semantically.
    - Optionally saves a plain-text preview for quick viewing.
    """
    import hashlib

    md_lines = []
    toc_lines = ["## 📚 Table of Contents\n"]

    with open(json_path, "r", encoding="utf-8") as f:
        blocks = json.load(f)

    for i, block in enumerate(blocks):
        block_type = block.get("type", "Block")
        text = block.get("text", "").strip()

        # Generate an anchor-friendly heading ID
        hash_id = hashlib.md5(text.encode("utf-8")).hexdigest()[:6]

        if block_type.lower() in ["title", "header"]:
            toc_lines.append(f"- [{text}](#{hash_id})")
            md_lines.append(f"### <a name=\"{hash_id}\"></a>{text}\n")

        elif block_type.lower() in ["listitem", "unorderedlist", "orderedlist"]:
            items = text.split("\n")
            for item in items:
                md_lines.append(f"- {item.strip()}")
            md_lines.append("")  # Spacer

        elif block_type.lower() in ["table", "tabular"]:
            # Handle as Markdown table
            rows = [line for line in text.split("\n") if line.strip()]
            if len(rows) >= 2:
                header = rows[0].split()
                separator = ["---"] * len(header)
                md_lines.append("| " + " | ".join(header) + " |")
                md_lines.append("| " + " | ".join(separator) + " |")
                for row in rows[1:]:
                    cols = row.split()
                    md_lines.append("| " + " | ".join(cols) + " |")
            else:
                md_lines.append(f"```table\n{text}\n```")

        elif block_type.lower() == "narrativetext":
            md_lines.append(f"> {text}")

        else:
            md_lines.append(f"**{block_type}**\n\n{text}")

        md_lines.append("\n---\n")

    toc_text = "\n".join(toc_lines) + "\n---\n"
    full_md = toc_text + "\n".join(md_lines)

    md_path = json_path.replace(".json", ".md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(full_md)

    # Optional: Plain-text preview (no Markdown formatting)
    if generate_preview:
        preview_text = "\n".join([block.get("text", "").strip() for block in blocks])
        preview_path = json_path.replace(".json", "_preview.txt")
        with open(preview_path, "w", encoding="utf-8") as f:
            f.write(preview_text)

    return md_path


def export_to_docx(json_path):
    doc = DocxDocument()
    with open(json_path, "r", encoding="utf-8") as f:
        blocks = json.load(f)

    for block in blocks:
        doc.add_heading(block.get("type", "Block"), level=2)
        doc.add_paragraph(block.get("text", ""))

    out_path = json_path.replace(".json", ".docx")
    doc.save(out_path)
    return out_path


def compute_global_anonymization_stats(
    client_name=None,
    project_id=None,
    service_id=None,
    date_from=None,
    date_to=None
):
    """
    Computes anonymization stats, optionally filtered by client/project/service/date.
    """

    from django.db.models import Q

    filters = Q(is_active=True)

    if client_name:
        filters &= Q(run__client_name=client_name)
    if project_id:
        filters &= Q(run__project_id=project_id)
    if service_id:
        filters &= Q(run__service_id=service_id)
    if date_from:
        filters &= Q(updated_at__date__gte=date_from)
    if date_to:
        filters &= Q(updated_at__date__lte=date_to)

    queryset = Anonymize.objects.filter(filters)

    files_with_entities = 0
    files_without_entities = 0
    total_entities_anonymized = 0
    entity_type_breakdown = {}

    for record in queryset:
        combined_map = {}
        if record.presidio_masking_map:
            combined_map.update(record.presidio_masking_map)
        if record.spacy_masking_map:
            combined_map.update(record.spacy_masking_map)

        entity_counts = {}
        for mask, original in combined_map.items():
            entity_type = mask.split("_MASKED_")[0]
            entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1

        file_entity_total = sum(entity_counts.values())
        total_entities_anonymized += file_entity_total

        for entity_type, count in entity_counts.items():
            entity_type_breakdown[entity_type] = (
                entity_type_breakdown.get(entity_type, 0) + count
            )

        if file_entity_total > 0:
            files_with_entities += 1
        else:
            files_without_entities += 1

    result = {
        "files_with_entities": files_with_entities,
        "files_without_entities": files_without_entities,
        "total_entities_anonymized": total_entities_anonymized,
        "entity_type_breakdown": entity_type_breakdown,
    }

    return result


def calculate_anonymization_insights(user):
    insights = {}

    # Runs by status
    runs = AnonymizationRun.objects.filter(client_name=user.username or user.email)
    status_counts = runs.values("status").annotate(count=Count("id"))
    insights["anonymization_status_counts"] = list(status_counts)

    # Average duration
    durations = []
    for run in runs:
        if run.updated_at and run.created_at:
            delta = (run.updated_at - run.created_at).total_seconds()
            durations.append(delta)
    insights["average_anonymization_time_seconds"] = sum(durations)/len(durations) if durations else 0

    # Risk levels
    risk_counts = (
        Anonymize.objects.filter(original_file__user=user)
        .values("risk_level")
        .annotate(count=Count("id"))
    )
    insights["risk_level_distribution"] = list(risk_counts)

    # High-risk documents
    high_risk_files = (
        Anonymize.objects
        .filter(original_file__user=user, risk_level="High")
        .values("original_file__filename", "risk_score")
        .order_by("-risk_score")[:10]
    )
    insights["top_high_risk_files"] = list(high_risk_files)

    # De-anonymization count
    deanonymize_count = (
        DeAnonymize.objects.filter(file__user=user).count()
    )
    insights["deanonymization_count"] = deanonymize_count

    return insights



# Reuse one heavy service (spaCy + Presidio) per worker
@lru_cache(maxsize=1)
def get_shared_anonymization_service():
    return AnonymizationService()

def list_supported_entities_runtime():
    """
    Returns the actual label sets available at runtime.
    Adds model/pipeline metadata, and avoids re-initializing heavy models.
    """
    data = {
        "spacy": [],
        "presidio": [],
        "meta": {"spacy_model": None, "spacy_pipes": []}
    }
    try:
        svc = get_shared_anonymization_service()

        # spaCy
        data["meta"]["spacy_model"] = getattr(getattr(svc, "nlp", None), "meta", {}).get("name")
        data["meta"]["spacy_pipes"] = list(getattr(svc.nlp, "pipe_names", []))
        if "ner" in getattr(svc.nlp, "pipe_names", []):
            data["spacy"] = sorted(set(svc.nlp.get_pipe("ner").labels))

        # Presidio
        data["presidio"] = sorted(set(svc.analyzer.get_supported_entities()))
    except Exception as e:
        logger.warning(f"[entities] Introspection failed: {e}")

    return data

# --- Optional: expose the risk weights the scorer uses (handy for docs/UI) ---
def risk_weight_for_label(label: str) -> float:
    """
    Keep this logic in sync with AnonymizationService.compute_risk_score().
    Returns the numeric weight used for a given (possibly non-canonical) label.
    """
    import re
    L = (label or "").upper().strip()
    # quick aliasing
    ALIASES = {"PER": "PERSON", "MISC": "MISC", "LOC": "LOC", "ORG": "ORG"}
    L = ALIASES.get(L, L)

    HIGH = {
        # financial & strong IDs
        "CREDIT_CARD","US_BANK_NUMBER","IBAN_CODE","CRYPTO",
        # national IDs / passports / drivers / tax etc. (from your list)
        "US_SSN","US_ITIN","US_PASSPORT","US_DRIVER_LICENSE",
        "UK_NINO","UK_NHS",
        "ES_NIF","ES_NIE",
        "IT_FISCAL_CODE","IT_DRIVER_LICENSE","IT_VAT_CODE","IT_PASSPORT","IT_IDENTITY_CARD",
        "PL_PESEL",
        "SG_NRIC_FIN","SG_UEN",
        "AU_ABN","AU_ACN","AU_TFN","AU_MEDICARE",
        "IN_PAN","IN_AADHAAR","IN_VEHICLE_REGISTRATION","IN_VOTER","IN_PASSPORT",
        "FI_PERSONAL_IDENTITY_CODE",
        "KR_RRN",
        "TH_TNIN",
    }
    MED_HIGH = {"MEDICAL_LICENSE","UK_NHS","AU_MEDICARE","SG_UEN","AU_ABN","AU_ACN","IT_VAT_CODE"}
    MED = {"EMAIL_ADDRESS","PHONE_NUMBER","IP_ADDRESS","URL","NRP","NORP"}
    LOW = {
        "PERSON","GPE","LOC","ORG","FAC","PRODUCT","EVENT","WORK_OF_ART",
        "LAW","LANGUAGE","DATE","TIME","PERCENT","MONEY","QUANTITY",
        "ORDINAL","CARDINAL","MISC","LOCATION","DATE_TIME"
    }

    if L in HIGH: return 8.0
    if L in MED_HIGH: return 6.0
    if L in MED: return 4.5
    if L in LOW: return 2.5

    # fallbacks for unknown/custom labels
    if re.search(r"(SSN|PASSPORT|DRIVER|LICENSE|NINO|NRIC|PESEL|TFN|AADHAAR|RRN|TNIN|IBAN|BANK|ACCOUNT|CREDIT|CARD|CRYPTO)", L):
        return 7.0
    if re.search(r"(EMAIL|PHONE|MOBILE|TEL|IP|URL|DOMAIN)", L):
        return 4.0
    if re.search(r"(DATE|TIME|MONEY|PERCENT|QUANTITY|ORDINAL|CARDINAL)", L):
        return 2.0
    return 3.0

def list_supported_entities_with_weights():
    """
    Returns { label: weight } for all labels currently supported by either engine.
    """
    data = list_supported_entities_runtime()
    labels = sorted(set(data.get("spacy", [])) | set(data.get("presidio", [])))
    return {lbl: risk_weight_for_label(lbl) for lbl in labels}

