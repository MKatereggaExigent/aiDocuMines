# core/utils.py
import os
import re
import csv
import logging
import mimetypes
import zipfile
import hashlib
import subprocess
from pathlib import Path
from datetime import datetime, timedelta
import xml.etree.ElementTree as ET

from django.conf import settings
from django.utils import timezone
from django.contrib.contenttypes.models import ContentType

# Third-party libs
from docx import Document
import PyPDF2
import fitz  # PyMuPDF
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from unidecode import unidecode
from tika import parser as tika_parser  # keep as used in extract_document_text
import pandas as pd

from core.models import File, Storage
from document_operations.models import Folder, FileFolderLink
from document_operations.utils import register_file_folder_link

logger = logging.getLogger(__name__)

# ✅ Allowed file types (MIME types)
ALLOWED_FILE_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "csv": "text/csv",
    "txt": "text/plain",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "tiff": "image/tiff",
    "gif": "image/gif",
}

# ✅ Max file size (100MB limit)
MAX_FILE_SIZE_MB = 100
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


def str_to_bool(value):
    """Converts various truthy/falsy string values to boolean."""
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    return str(value).strip().lower() in {"true", "yes", "1", "y", "t"}


def sanitize_filename(filename: str) -> str:
    """Sanitize the filename by removing/normalizing special characters."""
    # Replace German/special characters with ASCII
    filename = unidecode(filename)
    # Replace spaces
    filename = filename.replace(" ", "_")
    # Keep only safe chars
    filename = re.sub(r"[^a-zA-Z0-9\-_.]", "", filename)
    return filename


def rename_files_in_folder(folder_path: str) -> None:
    """Renames all files in a folder and subfolders by sanitizing filenames."""
    for root, _, files in os.walk(folder_path):
        for filename in files:
            src = os.path.join(root, filename)
            sanitized = sanitize_filename(filename)
            dst = os.path.join(root, sanitized)
            if src != dst:
                os.rename(src, dst)
                logger.info("Renamed %s -> %s", src, dst)


def save_uploaded_file(uploaded_file, storage_path: str, custom_filename=None):
    """
    Saves uploaded files synchronously.

    :param uploaded_file: File object from Django (InMemoryUploadedFile or TemporaryUploadedFile)
    :param storage_path: System location of uploaded files
    :param custom_filename: Optional new filename (sanitized automatically)
    :return: File metadata dictionary
    """
    # ✅ Validate file size
    if uploaded_file.size > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"File size exceeds {MAX_FILE_SIZE_MB}MB limit.")

    # ✅ Validate file type
    file_name = custom_filename or uploaded_file.name
    file_name = sanitize_filename(file_name)
    file_ext = file_name.split(".")[-1].lower()
    mime_type = mimetypes.guess_type(file_name)[0] or "application/octet-stream"

    if file_ext not in ALLOWED_FILE_TYPES:
        raise ValueError(f"Unsupported file type: {file_ext}")

    os.makedirs(storage_path, exist_ok=True)
    file_path = os.path.join(storage_path, file_name)

    # ✅ Save file in chunks
    chunk_size = 64 * 1024  # 64KB chunks
    with open(file_path, "wb") as out_file:
        while True:
            chunk = uploaded_file.read(chunk_size)
            if not chunk:
                break
            out_file.write(chunk)

    # Ensure all names in folder are sanitized (idempotent)
    rename_files_in_folder(storage_path)

    # ✅ Calculate file hash
    file_hash = calculate_md5(file_path)
    logger.info("✅ File %s saved at %s", file_name, file_path)

    return {
        "file_path": file_path,
        "filename": file_name,
        "file_size": uploaded_file.size,
        "file_type": mime_type,
        "md5_hash": file_hash,
        "upload_timestamp": datetime.utcnow().isoformat(),
    }


def calculate_md5(file_path: str) -> str:
    """Compute MD5 hash of a file."""
    md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            md5.update(chunk)
    return md5.hexdigest()


def convert_pdf_date(raw_date: str | None):
    """
    Converts extracted PDF (or ISO-ish) dates to ISO 8601 (YYYY-MM-DDTHH:MM:SS).
    Handles:
      - PDF format: "D:20250302161655+00'00'"
      - Generic/ISO-ish strings (best-effort)
    """
    if not raw_date or not isinstance(raw_date, str):
        return None

    # ✅ Handle PDF Standard Date Format
    m = re.match(r"D:(\d{14})([+-]\d{2})'(\d{2})'", raw_date)
    if m:
        date_part = m.group(1)  # "YYYYMMDDHHMMSS"
        tz_h = int(m.group(2))
        tz_m = int(m.group(3))
        try:
            parsed = datetime.strptime(date_part, "%Y%m%d%H%M%S")
            tz = timedelta(hours=tz_h, minutes=tz_m)
            parsed = parsed - tz
            return parsed.isoformat()
        except Exception:
            return None

    # ✅ Fallback: try parsing as generic date
    try:
        # Avoid pulling in dateutil parser globally again
        from dateutil import parser as dateutil_parser  # local import
        return dateutil_parser.parse(raw_date).isoformat()
    except Exception:
        return None


def extract_pdf_metadata(pdf_path: str) -> dict:
    """
    Extract extensive metadata from a PDF using multiple libraries.
    """
    if not os.path.exists(pdf_path):
        return {"Error": "File not found"}

    metadata: dict[str, object] = {}

    # ✅ PyMuPDF (fitz)
    try:
        doc = fitz.open(pdf_path)
        raw = doc.metadata or {}
        metadata.update({k.lower(): v for k, v in raw.items()})
        metadata["file_size"] = os.path.getsize(pdf_path)
        metadata["page_count"] = len(doc)
        metadata["creationdate"] = convert_pdf_date(metadata.get("creationdate"))
        metadata["moddate"] = convert_pdf_date(metadata.get("moddate"))
        metadata["is_encrypted"] = str_to_bool(raw.get("is_encrypted", "no"))
        metadata["encrypted"] = str_to_bool(raw.get("encrypted", "no"))
        metadata["optimized"] = str_to_bool(raw.get("optimized", "no"))
        metadata["pdf_version"] = raw.get("format", "Unknown PDF Version")

        # Fonts
        fonts = set()
        for page in doc:
            for font in page.get_fonts(full=True):
                # tuple layout may vary; index 3 is usually the font name
                if len(font) > 3 and font[3]:
                    fonts.add(font[3])
        metadata["fonts"] = sorted(fonts) if fonts else []
    except Exception as e:
        metadata.setdefault("errors", []).append(f"PyMuPDF error: {e}")

    # ✅ PyPDF2
    try:
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            doc_info = reader.metadata or {}
            metadata.update({k.lower(): v for k, v in doc_info.items()})
            try:
                metadata["page_rotation"] = [page.rotation for page in reader.pages]
            except Exception:
                pass

            # normalize common booleans
            for key in [
                "is_encrypted", "encrypted", "optimized", "tagged",
                "userproperties", "suspects", "custom_metadata"
            ]:
                if key in metadata:
                    metadata[key] = str_to_bool(metadata[key])
    except Exception as e:
        metadata.setdefault("errors", []).append(f"PyPDF2 error: {e}")

    # ✅ pdfminer.six
    try:
        with open(pdf_path, "rb") as f:
            p = PDFParser(f)
            d = PDFDocument(p)
            if getattr(d, "info", None):
                # d.info is typically a list of dictionaries
                info0 = d.info[0] if d.info else {}
                info_norm = {
                    (k.decode("utf-8", "ignore") if isinstance(k, bytes) else str(k)).lower():
                    (v.decode("utf-8", "ignore") if isinstance(v, bytes) else v)
                    for k, v in info0.items()
                }
                metadata.update(info_norm)
                for key in [
                    "is_encrypted", "encrypted", "optimized", "tagged",
                    "userproperties", "suspects", "custom_metadata"
                ]:
                    if key in metadata:
                        metadata[key] = str_to_bool(metadata[key])
    except Exception as e:
        metadata.setdefault("errors", []).append(f"pdfminer error: {e}")

    # ✅ pdfinfo (Poppler)
    try:
        out = subprocess.run(["pdfinfo", pdf_path], capture_output=True, text=True)
        lines = out.stdout.strip().splitlines()
        for line in lines:
            if ":" in line:
                k, v = line.split(":", 1)
                k = k.strip().lower().replace(" ", "_")
                metadata[k] = v.strip()
    except Exception as e:
        metadata.setdefault("errors", []).append(f"pdfinfo error: {e}")

    return metadata


def _read_docx_core_props(docx_zip: zipfile.ZipFile) -> dict:
    """Read docProps/core.xml; return normalized fields."""
    core = {}
    if "docProps/core.xml" not in docx_zip.namelist():
        return core
    root = ET.fromstring(docx_zip.read("docProps/core.xml").decode("utf-8"))
    ns = {
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
        "dc": "http://purl.org/dc/elements/1.1/",
        "dcterms": "http://purl.org/dc/terms/",
    }
    core["title"] = root.findtext("dc:title", default=None, namespaces=ns)
    core["creator"] = root.findtext("dc:creator", default=None, namespaces=ns)
    core["author"] = core.get("creator")  # mirror for convenience
    core["subject"] = root.findtext("dc:subject", default=None, namespaces=ns)
    core["keywords"] = root.findtext("cp:keywords", default=None, namespaces=ns)
    core["last_modified_by"] = root.findtext("cp:lastModifiedBy", default=None, namespaces=ns)
    core["category"] = root.findtext("cp:category", default=None, namespaces=ns)
    core["content_status"] = root.findtext("cp:contentStatus", default=None, namespaces=ns)
    core["revision"] = root.findtext("cp:revision", default=None, namespaces=ns)

    created_node = root.find("dcterms:created", ns)
    modified_node = root.find("dcterms:modified", ns)
    core["creationdate"] = convert_pdf_date(created_node.text) if (created_node is not None and created_node.text) else None
    core["moddate"] = convert_pdf_date(modified_node.text) if (modified_node is not None and modified_node.text) else None
    return core


def _read_docx_custom_props(docx_zip: zipfile.ZipFile) -> dict:
    """Read docProps/custom.xml; return {} if missing."""
    try:
        with docx_zip.open("docProps/custom.xml") as f:
            tree = ET.parse(f)
    except KeyError:
        return {}
    except Exception:
        return {}

    root = tree.getroot()
    ns = {
        "cp": "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties",
        "vt": "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes",
    }
    props = {}
    for prop in root.findall("cp:property", ns):
        name = prop.get("name")
        children = list(prop)
        if not children:
            continue
        value_elem = children[0]
        value = value_elem.text
        props[name] = value
    return props


def extract_docx_metadata(docx_path: str) -> dict:
    """
    DOCX metadata extraction. Always returns a dict that includes 'custom_metadata'.
    """
    metadata = {
        "format": "DOCX",
        "title": None,
        "author": None,
        "subject": None,
        "keywords": None,
        "creator": None,
        "producer": None,
        "creationdate": None,
        "moddate": None,
        "last_modified_by": None,
        "category": None,
        "content_status": None,
        "revision": None,
        "trapped": None,
        "encryption": "No",
        "file_size": os.path.getsize(docx_path) if os.path.exists(docx_path) else None,
        "page_count": None,
        "is_encrypted": False,
        "fonts": [],
        "page_rotation": None,
        "pdfminer_info": None,
        "metadata_stream": None,
        "tagged": None,
        "userproperties": None,
        "suspects": None,
        "form": None,
        "javascript": None,
        "pages": None,
        "encrypted": False,
        "page_size": None,
        "optimized": False,
        "pdf_version": None,
        "word_count": None,
        # 🔐 Always present to avoid KeyError downstream:
        "custom_metadata": {},
    }

    # Extract core & custom props from the package
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            core = _read_docx_core_props(z)
            metadata.update(core)
            metadata["custom_metadata"] = _read_docx_custom_props(z)
    except Exception as e:
        # Keep going; attach error note but preserve key presence
        metadata.setdefault("errors", []).append(f"core/custom props error: {e}")

    # Word count, fonts, rough page estimate via python-docx
    try:
        doc = Document(docx_path)
        paragraphs = list(doc.paragraphs)
        metadata["word_count"] = sum(len(p.text.split()) for p in paragraphs)
        # Very rough page heuristic; better replaced with actual pagination if needed
        metadata["page_count"] = max(1, len(paragraphs) // 20) if paragraphs else 1
        fonts = set()
        for p in paragraphs:
            for run in p.runs:
                if run.font and run.font.name:
                    fonts.add(run.font.name)
        metadata["fonts"] = sorted(fonts) if fonts else []
    except Exception as e:
        metadata.setdefault("errors", []).append(f"python-docx error: {e}")

    # DOCX doesn't really expose encrypted/optimized flags like PDFs; keep defaults
    return metadata


def extract_metadata(file_instance):
    """
    Extracts metadata from a given File model instance.
    """
    file_path = file_instance.filepath
    logger.info("🔍 Checking file path: %s", file_path)

    if not Path(file_path).exists():
        logger.error("❌ File does not exist: %s", file_path)
        return None

    logger.info("✅ File found! Extracting metadata for %s", file_path)

    base = {
        "file_id": file_instance.id,
        "storage_id": file_instance.storage.storage_id if file_instance.storage else None,
        "file_size": os.path.getsize(file_path),
        "format": mimetypes.guess_type(file_path)[0] or "unknown",
        "md5_hash": file_instance.md5_hash or None,
        "created_at": datetime.utcnow().isoformat(),
        "updated_at": datetime.utcnow().isoformat(),
        "title": None,
        "author": None,
        "subject": None,
        "keywords": None,
        "creator": None,
        "producer": None,
        "creationdate": None,
        "moddate": None,
        "last_modified_by": None,
        "category": None,
        "content_status": None,
        "revision": None,
        "trapped": None,
        "encryption": "No",
        "page_count": None,
        "is_encrypted": False,
        "fonts": [],
        "page_rotation": None,
        "pdfminer_info": None,
        "metadata_stream": None,
        "tagged": None,
        "userproperties": None,
        "suspects": None,
        "form": None,
        "javascript": None,
        "pages": None,
        "encrypted": False,
        "page_size": None,
        "optimized": False,
        "pdf_version": None,
        "word_count": None,
        # Make sure this key is *always* present
        "custom_metadata": {},
    }

    name_lc = (file_instance.filename or "").lower()

    if name_lc.endswith(".pdf"):
        try:
            pdf_meta = extract_pdf_metadata(file_path)
            base.update(pdf_meta)
            logger.info("✅ Extracted PDF metadata for file_id=%s", file_instance.id)
        except Exception as e:
            logger.error("❌ PDF metadata extraction failed: %s", e)

    elif name_lc.endswith(".docx"):
        try:
            docx_meta = extract_docx_metadata(file_path)
            base.update(docx_meta)
            logger.info("✅ Extracted DOCX metadata for file_id=%s", file_instance.id)
        except Exception as e:
            logger.error("❌ DOCX metadata extraction failed: %s", e)

    elif name_lc.endswith(".txt"):
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as text_file:
                base["word_count"] = sum(len(line.split()) for line in text_file)
            logger.info("✅ Extracted TXT metadata for file_id=%s", file_instance.id)
        except Exception as e:
            logger.error("❌ TXT metadata extraction failed: %s", e)

    else:
        # leave base as-is; format likely handled elsewhere
        pass

    return base


def extract_datetime_folder_from_path(filepath: str, project_id: str = None, service_id: str = None) -> str:
    """
    Extract the datetime folder (YYYYMMDD format) from a file path.

    The expected path structure is:
    .../uploads/{client_id}/{user_id}/{project_id}/{service_id}/{YYYYMMDD}/...

    Args:
        filepath: The full file path
        project_id: Optional project_id to help locate the datetime folder
        service_id: Optional service_id to help locate the datetime folder

    Returns:
        The datetime folder string (YYYYMMDD) or today's date if not found
    """
    if not filepath:
        return datetime.now().strftime("%Y%m%d")

    try:
        parts = filepath.replace('\\', '/').split('/')

        # Strategy 1: Find service_id in path and look for datetime folder after it
        if service_id:
            for i, part in enumerate(parts):
                if part == service_id and i + 1 < len(parts):
                    potential_date = parts[i + 1]
                    if len(potential_date) == 8 and potential_date.isdigit():
                        return potential_date

        # Strategy 2: Look for any 8-digit folder that looks like a date
        for part in parts:
            if len(part) == 8 and part.isdigit():
                # Validate it's a reasonable date (year between 2020 and 2030)
                year = int(part[:4])
                if 2020 <= year <= 2030:
                    return part

    except Exception as e:
        logger.debug(f"Could not extract datetime folder from path '{filepath}': {e}")

    return datetime.now().strftime("%Y%m%d")


def get_datetime_folder_from_dd_run(dd_run) -> str:
    """
    Extract the datetime folder from a DueDiligenceRun by looking at its associated files.

    Args:
        dd_run: A DueDiligenceRun instance

    Returns:
        The datetime folder string (YYYYMMDD) or today's date if not found
    """
    try:
        # Try to get files from document classifications
        if hasattr(dd_run, 'document_classifications'):
            classifications = dd_run.document_classifications.select_related('file').first()
            if classifications and classifications.file:
                return extract_datetime_folder_from_path(
                    classifications.file.filepath,
                    service_id=classifications.file.service_id
                )

        # Try to get files from the run
        if hasattr(dd_run, 'run') and hasattr(dd_run.run, 'files'):
            first_file = dd_run.run.files.first()
            if first_file:
                return extract_datetime_folder_from_path(
                    first_file.filepath,
                    service_id=first_file.service_id
                )
    except Exception as e:
        logger.debug(f"Could not extract datetime folder from DD run: {e}")

    return datetime.now().strftime("%Y%m%d")


def register_generated_file(file_path, user, run, project_id, service_id, folder_name="generated"):
    """
    Registers any generated file (e.g., translated, anonymized) in the File table
    using a flexible GenericForeignKey to support any run type.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Cannot register missing file: {file_path}")

    # Calculate MD5
    with open(file_path, "rb") as f:
        md5_hash = hashlib.md5(f.read()).hexdigest()

    # ContentType for the given run instance
    content_type = ContentType.objects.get_for_model(run)

    # ✅ Create Storage with generic linkage
    storage = Storage.objects.create(
        user=user,
        content_type=content_type,
        object_id=run.pk,
        upload_storage_location=file_path,
        output_storage_location=None,
    )

    # De-dup by md5 for this user
    existing = File.objects.filter(md5_hash=md5_hash, user=user).first()
    if existing:
        return existing

    from core.models import Run as RunModel
    file_instance = File.objects.create(
        run=run if isinstance(run, RunModel) else None,
        storage=storage,
        filename=os.path.basename(file_path),
        filepath=file_path,
        file_size=os.path.getsize(file_path),
        file_type=mimetypes.guess_type(file_path)[0] or os.path.splitext(file_path)[1].lstrip("."),
        md5_hash=md5_hash,
        user=user,
        project_id=project_id,
        service_id=service_id,
    )

    # Create folder link based on file path structure
    # This creates the proper nested folder hierarchy (e.g., 20250101/translations/Spanish)
    # and links the file to the leaf folder
    register_file_folder_link(file_instance)

    return file_instance


def extract_document_text(path, mime_type=None) -> str:
    """
    Extract text from various file types.
    """
    if not os.path.exists(path):
        return ""

    ext = os.path.splitext(path)[-1].lower()
    try:
        if ext == ".pdf":
            text = ""
            doc = fitz.open(path)
            for page in doc:
                text += page.get_text()
            return text

        if ext == ".docx":
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)

        if ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        if ext == ".csv":
            df = pd.read_csv(path, nrows=1000)
            return df.to_string(index=False)

        if ext == ".xlsx":
            df = pd.read_excel(path, nrows=1000)
            return df.to_string(index=False)

        # Fallback for other formats
        parsed = tika_parser.from_file(path)
        return (parsed.get("content") or "").strip()

    except Exception as e:
        return f"Error extracting text: {e}"


# =============================================================================
# SERVICE REPORT GENERATION UTILITIES
# =============================================================================

import json
import uuid
from typing import Any, Dict, List, Optional, Union


def generate_service_report_html(
    service_name: str,
    service_id: str,
    vertical: str,
    response_data: Any,
    input_files: Optional[List[Dict]] = None,
    query: Optional[str] = None,
    execution_time_seconds: Optional[float] = None,
    additional_metadata: Optional[Dict] = None
) -> str:
    """
    Generate a professional HTML report for any service response.

    Args:
        service_name: Human-readable service name (e.g., "Deal Document Search")
        service_id: Service ID (e.g., "pe-semantic-search")
        vertical: Vertical name (e.g., "Private Equity", "Class Actions")
        response_data: The service response data (dict, list, or any serializable)
        input_files: Optional list of input files with {filename, file_id, path}
        query: Optional query string if this was a search/interrogation service
        execution_time_seconds: Optional execution time
        additional_metadata: Optional extra metadata to include

    Returns:
        HTML string of the report
    """
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Format response data for display
    if isinstance(response_data, (dict, list)):
        formatted_response = json.dumps(response_data, indent=2, default=str)
    else:
        formatted_response = str(response_data)

    # Build input files section
    input_files_html = ""
    if input_files:
        input_files_html = """
        <div class="section">
            <h2>📁 Input Files</h2>
            <ul class="files-list">
        """
        for f in input_files:
            filename = f.get('filename') or f.get('name') or 'Unknown'
            file_id = f.get('file_id') or f.get('id') or ''
            input_files_html += f'<li><strong>{filename}</strong> (ID: {file_id})</li>'
        input_files_html += "</ul></div>"

    # Build query section
    query_html = ""
    if query:
        query_html = f"""
        <div class="section">
            <h2>🔍 Query</h2>
            <div class="query-box">{query}</div>
        </div>
        """

    # Build metadata section
    metadata_html = ""
    if additional_metadata:
        metadata_html = """
        <div class="section">
            <h2>📊 Additional Metadata</h2>
            <div class="metadata">
        """
        for key, value in additional_metadata.items():
            metadata_html += f"""
            <div class="metadata-item">
                <div class="metadata-label">{key.replace('_', ' ').title()}</div>
                <div class="metadata-value">{value}</div>
            </div>
            """
        metadata_html += "</div></div>"

    # Calculate result summary
    result_count = 0
    if isinstance(response_data, list):
        result_count = len(response_data)
    elif isinstance(response_data, dict):
        if 'results' in response_data:
            result_count = len(response_data.get('results', []))
        elif 'data' in response_data:
            data = response_data.get('data')
            result_count = len(data) if isinstance(data, list) else 1
        else:
            result_count = len(response_data)

    html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{service_name} - Service Report</title>
    <style>
        :root {{
            --primary-color: #667eea;
            --secondary-color: #764ba2;
            --bg-color: #f8fafc;
            --card-bg: #ffffff;
            --text-color: #2d3748;
            --text-muted: #718096;
            --border-color: #e2e8f0;
            --success-color: #48bb78;
            --warning-color: #ed8936;
            --danger-color: #f56565;
        }}

        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, 'Roboto', 'Helvetica Neue', sans-serif;
            background: var(--bg-color);
            color: var(--text-color);
            line-height: 1.6;
            padding: 20px;
        }}

        .container {{
            max-width: 1200px;
            margin: 0 auto;
        }}

        .header {{
            background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
            color: white;
            padding: 40px;
            border-radius: 16px;
            margin-bottom: 30px;
            box-shadow: 0 10px 40px rgba(102, 126, 234, 0.3);
        }}

        .header h1 {{
            font-size: 2rem;
            font-weight: 600;
            margin-bottom: 8px;
        }}

        .header .subtitle {{
            opacity: 0.9;
            font-size: 1rem;
        }}

        .header .meta {{
            display: flex;
            gap: 20px;
            margin-top: 20px;
            flex-wrap: wrap;
        }}

        .header .meta-item {{
            background: rgba(255,255,255,0.2);
            padding: 8px 16px;
            border-radius: 20px;
            font-size: 0.9rem;
        }}

        .section {{
            background: var(--card-bg);
            padding: 25px;
            margin-bottom: 20px;
            border-radius: 12px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
            border: 1px solid var(--border-color);
        }}

        .section h2 {{
            color: var(--text-color);
            font-size: 1.25rem;
            font-weight: 600;
            margin-bottom: 15px;
            padding-bottom: 10px;
            border-bottom: 2px solid var(--border-color);
        }}

        .metadata {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 15px;
        }}

        .metadata-item {{
            padding: 15px;
            background: var(--bg-color);
            border-radius: 8px;
            border-left: 4px solid var(--primary-color);
        }}

        .metadata-label {{
            font-weight: 600;
            color: var(--text-muted);
            font-size: 0.8rem;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin-bottom: 5px;
        }}

        .metadata-value {{
            color: var(--text-color);
            font-size: 1rem;
        }}

        .files-list {{
            list-style: none;
            padding: 0;
        }}

        .files-list li {{
            padding: 12px 16px;
            background: var(--bg-color);
            margin: 8px 0;
            border-radius: 8px;
            border-left: 4px solid var(--success-color);
        }}

        .query-box {{
            background: var(--bg-color);
            padding: 20px;
            border-radius: 8px;
            border: 1px solid var(--border-color);
            font-style: italic;
            color: var(--text-muted);
        }}

        .results-data {{
            background: #1a202c;
            color: #e2e8f0;
            padding: 20px;
            border-radius: 8px;
            overflow-x: auto;
            font-family: 'Monaco', 'Menlo', 'Ubuntu Mono', monospace;
            font-size: 0.85rem;
            line-height: 1.5;
            max-height: 600px;
            overflow-y: auto;
        }}

        .results-data pre {{
            margin: 0;
            white-space: pre-wrap;
            word-break: break-word;
        }}

        .summary-stats {{
            display: flex;
            gap: 20px;
            flex-wrap: wrap;
            margin-bottom: 20px;
        }}

        .stat-card {{
            flex: 1;
            min-width: 150px;
            padding: 20px;
            background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
            color: white;
            border-radius: 12px;
            text-align: center;
        }}

        .stat-card .stat-value {{
            font-size: 2rem;
            font-weight: 700;
        }}

        .stat-card .stat-label {{
            font-size: 0.85rem;
            opacity: 0.9;
        }}

        .footer {{
            text-align: center;
            padding: 20px;
            color: var(--text-muted);
            font-size: 0.85rem;
        }}

        @media print {{
            body {{
                background: white;
                padding: 0;
            }}
            .header {{
                box-shadow: none;
            }}
            .section {{
                box-shadow: none;
                break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📋 {service_name}</h1>
            <div class="subtitle">Service Report - {vertical}</div>
            <div class="meta">
                <span class="meta-item">📅 {timestamp}</span>
                <span class="meta-item">🔧 {service_id}</span>
                {f'<span class="meta-item">⏱️ {execution_time_seconds:.2f}s</span>' if execution_time_seconds else ''}
                <span class="meta-item">📊 {result_count} results</span>
            </div>
        </div>

        <div class="summary-stats">
            <div class="stat-card">
                <div class="stat-value">{result_count}</div>
                <div class="stat-label">Total Results</div>
            </div>
            {f'''<div class="stat-card">
                <div class="stat-value">{len(input_files) if input_files else 0}</div>
                <div class="stat-label">Input Files</div>
            </div>''' if input_files else ''}
            {f'''<div class="stat-card">
                <div class="stat-value">{execution_time_seconds:.1f}s</div>
                <div class="stat-label">Execution Time</div>
            </div>''' if execution_time_seconds else ''}
        </div>

        {query_html}

        {input_files_html}

        {metadata_html}

        <div class="section">
            <h2>📄 Results Data</h2>
            <div class="results-data">
                <pre>{formatted_response}</pre>
            </div>
        </div>

        <div class="footer">
            <p>Generated by aiDocuMines &bull; {vertical} &bull; {timestamp}</p>
        </div>
    </div>
</body>
</html>
    """

    return html


def generate_and_register_service_report(
    service_name: str,
    service_id: str,
    vertical: str,
    response_data: Any,
    user,
    run,
    project_id: str,
    service_id_folder: str,
    folder_name: Optional[str] = None,
    input_files: Optional[List[Dict]] = None,
    query: Optional[str] = None,
    execution_time_seconds: Optional[float] = None,
    additional_metadata: Optional[Dict] = None,
    file_format: str = "html",
    datetime_folder: Optional[str] = None
) -> Dict:
    """
    Generate a service report and register it in the file tree.

    The folder structure follows the pattern:
    {MEDIA_ROOT}/uploads/{client_id}/{user_id}/{project_id}/{service_id}/{datetime}/{output_folder}/

    Args:
        service_name: Human-readable service name
        service_id: Service ID
        vertical: Vertical name
        response_data: The service response data
        user: The user who executed the service
        run: The run instance (for GenericForeignKey)
        project_id: Project ID for folder organization
        service_id_folder: Service ID for folder organization
        folder_name: Custom folder name (default: derived from service_name)
        input_files: Optional list of input files (used to extract datetime_folder if not provided)
        query: Optional query string
        execution_time_seconds: Optional execution time
        additional_metadata: Optional extra metadata
        file_format: "html" or "json" (default: "html")
        datetime_folder: The datetime folder (YYYYMMDD format) - if not provided, extracted from input_files or uses today

    Returns:
        Dict with registered file info: {file_id, filename, filepath, folder_name}
    """
    from django.conf import settings

    # Generate folder name from service name if not provided
    if not folder_name:
        folder_name = service_name.lower().replace(' ', '-').replace('_', '-')

    # Determine datetime_folder from input files if not provided
    if not datetime_folder and input_files:
        for input_file in input_files:
            filepath = input_file.get('filepath') or input_file.get('file_path') or ''
            # Try to extract datetime folder from filepath pattern: .../{project_id}/{service_id}/{YYYYMMDD}/...
            if project_id and service_id_folder and filepath:
                try:
                    # Find the datetime folder in the path
                    parts = filepath.split(os.sep)
                    for i, part in enumerate(parts):
                        if part == service_id_folder and i + 1 < len(parts):
                            potential_date = parts[i + 1]
                            # Check if it looks like a date folder (YYYYMMDD format)
                            if len(potential_date) == 8 and potential_date.isdigit():
                                datetime_folder = potential_date
                                logger.debug(f"Extracted datetime_folder from input file: {datetime_folder}")
                                break
                except Exception as e:
                    logger.debug(f"Could not extract datetime from input file path: {e}")
            if datetime_folder:
                break

    # Default to today's date if still not determined
    if not datetime_folder:
        datetime_folder = datetime.now().strftime("%Y%m%d")

    # Get client_id
    client_id = getattr(user, 'client_id', None) or 'default'

    # Generate unique filename
    timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = str(uuid.uuid4())[:8]

    if file_format == "json":
        filename = f"{folder_name}_report_{timestamp_str}_{unique_id}.json"
        content = json.dumps(response_data, indent=2, default=str)
        content_type = "application/json"
    else:
        filename = f"{folder_name}_report_{timestamp_str}_{unique_id}.html"
        content = generate_service_report_html(
            service_name=service_name,
            service_id=service_id,
            vertical=vertical,
            response_data=response_data,
            input_files=input_files,
            query=query,
            execution_time_seconds=execution_time_seconds,
            additional_metadata=additional_metadata
        )
        content_type = "text/html"

    # Determine file path with correct structure:
    # {MEDIA_ROOT}/uploads/{client_id}/{user_id}/{project_id}/{service_id}/{datetime}/{output_folder}/
    base_path = os.path.join(
        settings.MEDIA_ROOT,
        "uploads",
        str(client_id),
        str(user.id),
        project_id,
        service_id_folder,
        datetime_folder,
        folder_name
    )
    os.makedirs(base_path, exist_ok=True)
    file_path = os.path.join(base_path, filename)

    # Write the file
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)

    # Register the file using existing utility
    registered_file = register_generated_file(
        file_path=file_path,
        user=user,
        run=run,
        project_id=project_id,
        service_id=service_id_folder,
        folder_name=folder_name
    )

    logger.info(f"✅ Generated and registered service report: {filename} in {project_id}/{service_id_folder}/{datetime_folder}/{folder_name}")

    return {
        "file_id": registered_file.id,
        "filename": registered_file.filename,
        "filepath": registered_file.filepath,
        "folder_name": folder_name,
        "datetime_folder": datetime_folder,
        "download_url": f"/api/v1/documents/download/{registered_file.id}/",
        "content_type": content_type
    }


def generate_search_results_report(
    query: str,
    results: List[Dict],
    user,
    run,
    project_id: str,
    service_id: str,
    search_type: str = "semantic",
    vertical: str = "AI Services",
    execution_time_seconds: Optional[float] = None
) -> Dict:
    """
    Convenience function specifically for search service results.

    Args:
        query: The search query
        results: List of search results
        user: The user who executed the search
        run: The run instance
        project_id: Project ID
        service_id: Service ID
        search_type: "semantic" or "elasticsearch"
        vertical: Vertical name
        execution_time_seconds: Optional execution time

    Returns:
        Dict with registered file info
    """
    service_name = f"{search_type.title()} Search"
    folder_name = f"{search_type}-search-results"

    return generate_and_register_service_report(
        service_name=service_name,
        service_id=f"ai-{search_type}-search",
        vertical=vertical,
        response_data={"query": query, "results": results, "total": len(results)},
        user=user,
        run=run,
        project_id=project_id,
        service_id_folder=service_id,
        folder_name=folder_name,
        query=query,
        execution_time_seconds=execution_time_seconds,
        additional_metadata={
            "search_type": search_type,
            "result_count": len(results)
        }
    )

