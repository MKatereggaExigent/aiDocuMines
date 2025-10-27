import os
import logging
import pandas as pd
import fitz
from celery import shared_task, chord
from django.db import transaction
from django.shortcuts import get_object_or_404
from django.utils.timezone import now
from core.models import File
import shutil
from docx import Document
from document_ocr.models import OCRRun, OCRFile
from document_ocr.utils import OCRService, cleanup_tmp_dir
import uuid
from core.utils import register_generated_file
# from glob import glob
import glob 
from PyPDF2 import PdfMerger

logger = logging.getLogger(__name__)

@shared_task
def process_ocr(run_id, file_id, ocr_option="basic"):
    """
    Celery task to perform OCR on a document asynchronously.
    """
    try:
        # Retrieve the OCRRun object and file
        run = OCRRun.objects.get(id=run_id)
        file_obj = get_object_or_404(File, id=file_id)

        # Create or get the OCRFile entry
        ocr_file, created = OCRFile.objects.get_or_create(
            original_file=file_obj, ocr_option=ocr_option, defaults={"status": "Processing", "run": run}
        )

        run.status = "Processing"
        run.save()

        ocr_service = OCRService()

        # Check if the file is a valid PDF
        if not ocr_service.is_pdf(file_obj.filepath):
            logger.info(f"🔹 Skipping OCR: {file_obj.filepath} is not a PDF.")
            ocr_file.status = "Completed"
            ocr_file.save()
            return {"message": "File is not a PDF. OCR skipped."}

        # Step 1: Extract bookmarks before processing
        bookmarks_df = ocr_service.extract_bookmarks_to_dataframe(file_obj.filepath)
        bookmarks_list = bookmarks_df.to_dict(orient="records")  # Convert to list for Celery serialization

        # Step 2: Burst the PDF into page batches
        batch_files = ocr_service.burst_pdf(file_obj.filepath, ocr_option=ocr_option)

        if not batch_files:
            ocr_file.status = "Failed"
            ocr_file.save()
            return {"error": "Failed to burst the PDF", "file_path": file_obj.filepath}

        # Ensure output directories exist
        output_dir = os.path.join(os.path.dirname(file_obj.filepath), "ocr", ocr_option.lower(), "tmp")
        os.makedirs(output_dir, exist_ok=True)

        # Step 3: Define OCR tasks for each batch
        ocr_tasks = [
            ocr_pdf_page_batch.s(ocr_file.id, batch_file, start_page, end_page, ocr_option)  # Pass ocr_option
            for start_page, end_page, batch_file in batch_files
        ]

        # Step 4: Wait for all OCR tasks to finish before merging
        workflow = chord(ocr_tasks)(merge_ocr_batches.s(ocr_file.id, bookmarks_list))  # Pass bookmarks_list

        return workflow

    except Exception as e:
        logger.error(f"❌ OCR processing failed: {e}")
        if 'ocr_file' in locals():
            ocr_file.status = "Failed"
            ocr_file.save()
        return {"error": str(e)}

@shared_task
def ocr_pdf_page_batch(file_id, batch_file_path, start_page, end_page, ocr_option="basic"):
    """OCR task for a batch of pages in a PDF."""
    try:
        # Ensure the file exists
        batch_file_path = os.path.normpath(batch_file_path)  # Fix path issues
        if not os.path.exists(batch_file_path):
            logger.error(f"❌ Batch file not found: {batch_file_path}")
            return {"error": "Batch file not found", "batch_file_path": batch_file_path}

        ocr_service = OCRService()

        # Apply OCR to the batch file
        ocr_file = ocr_service.apply_ocr(file_id, batch_file_path, ocr_option)  # Pass ocr_option
        return {"start_page": start_page, "end_page": end_page, "ocr_file": ocr_file}

    except Exception as e:
        logger.error(f"❌ OCR processing failed: {str(e)}")
        return {"error": str(e), "batch_file_path": batch_file_path}




@shared_task
def merge_ocr_batches(results, ocr_file_id, bookmarks_list):
    """Merges all OCR'ed PDF batches and reattaches bookmarks."""
    file_entry = get_object_or_404(OCRFile, id=ocr_file_id)

    # Ensure the OCR option is passed correctly
    if not file_entry.ocr_option:
        logger.error(f"❌ OCR option is missing for OCRFile {ocr_file_id}.")
        file_entry.status = "Failed"
        file_entry.save()
        return {"error": "OCR option missing"}

    ocr_service = OCRService()

    # Locate the tmp directory where the batch PDFs are stored
    ocr_dir = os.path.join(os.path.dirname(file_entry.original_file.filepath), "ocr", file_entry.ocr_option.lower(), "tmp")
    logger.info(f"🔍 Searching for OCR tmp directory: {ocr_dir}"
                )
    if not os.path.exists(ocr_dir):
        if "basic-ocr" in ocr_dir:
            ocr_dir = ocr_dir.replace("basic-ocr", "advanced-ocr")
            
            logger.info(f"🔄 Updated OCR tmp directory to: {ocr_dir}")
            
            # There is a folder inside the ocr_dir that needs to be removed. i.e delete os.path.join(ocr_dir, "ocr")
            
            tmp_ocr_dir = os.path.join(os.path.dirname(file_entry.original_file.filepath), "ocr", file_entry.ocr_option.lower(), "ocr")
            tmp_ocr_dir = tmp_ocr_dir.replace("basic-ocr", "advanced-ocr")
            
            if os.path.exists(tmp_ocr_dir):
                # Remove the tmp_ocr_dir
                shutil.rmtree(tmp_ocr_dir)
                logger.info(f"🔄 Removed OCR folder inside tmp directory: {tmp_ocr_dir}")
                
        else:
            logger.error(f"❌ OCR tmp directory not found: {ocr_dir}.")
             
    
    # Find all PDF files under the OCR tmp directory (recursively)
    pdf_files = glob.glob(os.path.join(ocr_dir, '*.pdf'))

    # Ensure there are files to merge
    if not pdf_files:
        logger.error(f"❌ No OCR batches found to merge in {ocr_dir}.")
        file_entry.status = "Failed"
        file_entry.save()
        return {"error": "No OCR batches found"}

    logger.info(f"🔄 Found {len(pdf_files)} PDFs to merge.")

    # Merge PDFs using PdfMerger
    final_pdf_dir = os.path.join(os.path.dirname(file_entry.original_file.filepath), "ocr", file_entry.ocr_option.lower(), "final")
    if "advanced-ocr" in ocr_dir:
        final_pdf_dir = final_pdf_dir.replace("basic-ocr", "advanced-ocr")
        
    logger.info(f"🔄 Final PDF directory: {final_pdf_dir}")
    
    os.makedirs(final_pdf_dir, exist_ok=True)

    try:
        final_pdf_path = os.path.join(final_pdf_dir, f"ocr-{uuid.uuid4()}.pdf")
        try:
            merger = PdfMerger()
            for pdf in pdf_files:
                logger.info(f"🔹 Adding {pdf} to the merger.")
                merger.append(pdf)

            # Write the merged PDF to the final output path
            with open(final_pdf_path, "wb") as output_file:
                merger.write(output_file)

            logger.info(f"✅ Merged PDF saved to: {final_pdf_path}")
        except Exception as e:
            logger.error(f"❌ Error during merging PDFs: {e}")
            file_entry.status = "Failed"
            file_entry.save()
            return {"error": f"Error merging PDFs: {e}"}

        # Register the OCR'ed PDF
        upload_run = file_entry.original_file.run
        registered = register_generated_file(
            file_path=final_pdf_path,
            user=file_entry.original_file.user,
            run=upload_run,
            project_id=file_entry.original_file.project_id,
            service_id=file_entry.original_file.service_id,
            folder_name="ocr"
        )

        # Reattach bookmarks
        pdf_document = fitz.open(final_pdf_path)
        total_pages = pdf_document.page_count
        pdf_document.close()

        bookmarks_df = pd.DataFrame(bookmarks_list)

        try:
            final_pdf_with_bookmarks = ocr_service.reattach_bookmarks_from_dataframe(final_pdf_path, bookmarks_df, 1, total_pages)
        except Exception as e:
            logger.error(f"❌ Failed to reattach bookmarks: {e}")
            final_pdf_with_bookmarks = final_pdf_path


        # Register the final PDF with bookmarks
        registered_pdf = register_generated_file(
            file_path=final_pdf_with_bookmarks,  # Use the file with bookmarks
            user=file_entry.original_file.user,
            run=file_entry.original_file.run,
            project_id=file_entry.original_file.project_id,
            service_id=file_entry.original_file.service_id,
            folder_name="ocr"
        )

        # Update the database record
        with transaction.atomic():
            file_entry.ocr_filepath = final_pdf_path
            file_entry.status = "Processed"
            file_entry.updated_at = now()
            file_entry.save()

        # Cleanup temporary files
        ocr_service.cleanup_tmp_dir(ocr_dir)

        # Trigger DOCX conversion
        process_pdf_to_docx.delay(file_entry.id, final_pdf_path)

        logger.info(f"✅ OCR processing completed and saved: {final_pdf_path}")

    # If we enter the except block, this part of the code will handle the final file detection in Advanced-ocr case
    except Exception as e:
        logger.error(f"❌ Advanced OCR processing failed, attempting to detect final file: {e}")

        try:
            # Get the last final file found (i.e. final_merged_ocr_file.pdf) and re-attach bookmarks
            final_pdf_dir = os.path.join(os.path.dirname(file_entry.original_file.filepath), "ocr", file_entry.ocr_option.lower(), "final")
            files_in_final = glob.glob(os.path.join(final_pdf_dir, "*final_merged_ocr_file.pdf"))

            if files_in_final:
                final_pdf_with_bookmarks = files_in_final[0]
                logger.info(f"✅ Found final PDF at: {final_pdf_with_bookmarks}")

                # Register the final file
                registered_pdf = register_generated_file(
                    file_path=final_pdf_with_bookmarks,
                    user=file_entry.original_file.user,
                    run=file_entry.original_file.run,
                    project_id=file_entry.original_file.project_id,
                    service_id=file_entry.original_file.service_id,
                    folder_name="ocr"
                )

                # Update the database record
                with transaction.atomic():
                    file_entry.ocr_filepath = final_pdf_with_bookmarks
                    file_entry.status = "Processed"
                    file_entry.updated_at = now()
                    file_entry.save()

                logger.info(f"✅ Advanced OCR processed and saved: {final_pdf_with_bookmarks}")

                # Cleanup temporary files
                ocr_service.cleanup_tmp_dir(ocr_dir)
                return {"status": "Completed", "ocr_merged_pdf": final_pdf_with_bookmarks}
            else:
                raise FileNotFoundError("Final OCR PDF not found in the 'final' folder.")

        except Exception as e:
            logger.error(f"❌ Advanced OCR failed to detect final file: {e}")
            file_entry.status = "Failed"
            file_entry.save()
            return {"error": "Advanced OCR final file not found or error occurred."}


    return {
        "ocr_run_id": str(file_entry.run.id),
        "file_id": file_entry.original_file.id,
        "ocr_file_id": file_entry.id,
        "ocr_merged_pdf": final_pdf_path,
        "status": "Completed",
        "registered_outputs": [{
            "file_id": registered.id,
            "filename": registered.filename,
            "path": registered.filepath
        }]
    }



@shared_task
def process_pdf_to_docx(ocr_file_id, final_pdf_path):
    """Converts an OCR’ed PDF to formatted DOCX and registers it in the File model."""
    file_entry = get_object_or_404(OCRFile, id=ocr_file_id)

    # Ensure the OCRed PDF exists before conversion
    if not final_pdf_path or not os.path.exists(final_pdf_path):
        logger.error(f"❌ No OCRed PDF found at {final_pdf_path}. Skipping DOCX conversion.")
        return {"error": "OCRed PDF not found", "ocr_file_id": ocr_file_id}

    ocr_service = OCRService()
    formatted_docx_path = os.path.join(os.path.dirname(final_pdf_path), f"{uuid.uuid4()}_formatted.docx")

    # Convert to formatted DOCX
    formatted_output = ocr_service.convert_to_formatted_docx(final_pdf_path, formatted_docx_path)

    registered = None
    if formatted_output:
        with transaction.atomic():
            file_entry.docx_path = formatted_output
            file_entry.save()

        # Register in File table
        registered = register_generated_file(
            file_path=formatted_output,
            user=file_entry.original_file.user,
            run=file_entry.original_file.run,
            project_id=file_entry.original_file.project_id,
            service_id=file_entry.original_file.service_id,
            folder_name="ocr"
        )

        # Trigger raw DOCX generation
        generate_raw_docx.delay(file_entry.id)

    return {
        "ocr_run_id": str(file_entry.run.id),
        "file_id": file_entry.original_file.id,
        "ocr_file_id": file_entry.id,
        "formatted_docx": formatted_output,
        "status": "Completed" if formatted_output else "Failed",
        "registered_outputs": [{
            "file_id": registered.id,
            "filename": registered.filename,
            "path": registered.filepath
        }] if registered else []
    }

@shared_task
def generate_raw_docx(ocr_file_id):
    """Generate and register a raw DOCX with unformatted text extracted from the formatted DOCX."""
    file_entry = get_object_or_404(OCRFile, id=ocr_file_id)

    if not file_entry.docx_path or not os.path.exists(file_entry.docx_path):
        logger.error(f"❌ No formatted DOCX found for ocr_file_id: {ocr_file_id}. Skipping RAW DOCX creation.")
        return {
            "ocr_file_id": ocr_file_id,
            "status": "Failed",
            "error": "Formatted DOCX not found"
        }

    raw_docx_path = os.path.join(os.path.dirname(file_entry.docx_path), f"{uuid.uuid4()}_raw.docx")

    original_doc = Document(file_entry.docx_path)
    raw_doc = Document()

    for para in original_doc.paragraphs:
        if para.text.strip():
            new_para = raw_doc.add_paragraph(para.text)

            new_para.paragraph_format.left_indent = None
            new_para.paragraph_format.right_indent = None
            new_para.paragraph_format.first_line_indent = None
            new_para.paragraph_format.space_before = None
            new_para.paragraph_format.space_after = None
            new_para.paragraph_format.alignment = None

            for run in new_para.runs:
                run.bold = False
                run.italic = False
                run.underline = False

    raw_doc.save(raw_docx_path)

    # Register file
    registered = register_generated_file(
        file_path=raw_docx_path,
        user=file_entry.original_file.user,
        run=file_entry.original_file.run,
        project_id=file_entry.original_file.project_id,
        service_id=file_entry.original_file.service_id,
        folder_name="ocr"
    )

    # Save to DB
    with transaction.atomic():
        file_entry.raw_docx_path = raw_docx_path
        file_entry.save()

    logger.info(f"✅ Raw DOCX created and registered: {raw_docx_path}")

    return {
        "ocr_run_id": str(file_entry.run.id),
        "file_id": file_entry.original_file.id,
        "ocr_file_id": file_entry.id,
        "raw_docx": raw_docx_path,
        "status": "Completed",
        "registered_outputs": [{
            "file_id": registered.id,
            "filename": registered.filename,
            "path": registered.filepath
        }]
    }

