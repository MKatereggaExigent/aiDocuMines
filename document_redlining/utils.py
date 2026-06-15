import os
import difflib
import logging
import shutil
import subprocess
from difflib import SequenceMatcher
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from copy import deepcopy

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path):
    """Extract text content from a file based on its extension."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[-1].lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if ext == ".docx":
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(file_path)
            text = "".join(page.get_text() for page in doc)
            doc.close()
            return text
        except ImportError:
            try:
                from pdfminer.high_level import extract_text
                return extract_text(file_path)
            except ImportError:
                raise ImportError("No PDF extraction library available (install PyMuPDF or pdfminer.six)")

    if ext == ".html" or ext == ".htm":
        try:
            from bs4 import BeautifulSoup
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                return soup.get_text()
        except ImportError:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

    raise ValueError(f"Unsupported file type: {ext}")


def _add_red_strikethrough_run(p_elem, text):
    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FF0000')
    rPr.append(color)
    strike = OxmlElement('w:strike')
    strike.set(qn('w:val'), '1')
    rPr.append(strike)
    run_elem.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    run_elem.append(t_elem)
    p_elem.append(run_elem)
    return run_elem


def _add_blue_underline_run(p_elem, text):
    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    u_elem = OxmlElement('w:u')
    u_elem.set(qn('w:val'), 'single')
    rPr.append(u_elem)
    run_elem.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    run_elem.append(t_elem)
    p_elem.append(run_elem)
    return run_elem


def _add_normal_run(p_elem, text):
    run_elem = OxmlElement('w:r')
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    run_elem.append(t_elem)
    p_elem.append(run_elem)
    return run_elem


def _copy_pPr(orig_para, p_elem):
    pPr = orig_para._element.find(qn('w:pPr'))
    if pPr is not None:
        p_elem.insert(0, deepcopy(pPr))


def _create_p_element(body, orig_para=None):
    p_elem = OxmlElement('w:p')
    body.append(p_elem)
    if orig_para is not None:
        _copy_pPr(orig_para, p_elem)
    return p_elem


def _build_modified_paragraph(body, orig_para, comp_para, author):
    p_elem = _create_p_element(body, orig_para or comp_para)

    text1 = orig_para.text if orig_para else ''
    text2 = comp_para.text if comp_para else ''

    if not text1 and not text2:
        return p_elem

    matcher = SequenceMatcher(None, text1, text2)
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'equal':
            _add_normal_run(p_elem, text1[i1:i2])
        elif op == 'replace':
            if i1 < i2:
                _add_red_strikethrough_run(p_elem, text1[i1:i2])
            if j1 < j2:
                _add_blue_underline_run(p_elem, text2[j1:j2])
        elif op == 'delete':
            if i1 < i2:
                _add_red_strikethrough_run(p_elem, text1[i1:i2])
        elif op == 'insert':
            if j1 < j2:
                _add_blue_underline_run(p_elem, text2[j1:j2])

    return p_elem


def _build_equal_paragraph(body, orig_para):
    p_elem = _create_p_element(body, orig_para)
    _add_normal_run(p_elem, orig_para.text)
    return p_elem


def _build_deleted_paragraph(body, orig_para):
    p_elem = _create_p_element(body, orig_para)
    _add_red_strikethrough_run(p_elem, orig_para.text)
    return p_elem


def _build_added_paragraph(body, comp_para):
    p_elem = _create_p_element(body, comp_para)
    _add_blue_underline_run(p_elem, comp_para.text)
    return p_elem


def generate_redline_docx(paragraph_pairs, original_path, output_path, author):
    doc = Document(original_path)

    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)

    for pair in paragraph_pairs:
        ptype = pair['type']
        if ptype == 'equal' and pair.get('original'):
            _build_equal_paragraph(body, pair['original'])
        elif ptype == 'modified':
            _build_modified_paragraph(body, pair.get('original'), pair.get('comparison'), author)
        elif ptype == 'deleted' and pair.get('original'):
            _build_deleted_paragraph(body, pair['original'])
        elif ptype == 'added' and pair.get('comparison'):
            _build_added_paragraph(body, pair['comparison'])

    doc.save(output_path)
    logger.info(f"Redline DOCX saved to {output_path}")
    return output_path


def compare_docx_files(docx1_path, docx2_path, author="Redlining System"):
    doc1 = Document(docx1_path)
    doc2 = Document(docx2_path)

    paras1 = doc1.paragraphs
    paras2 = doc2.paragraphs
    texts1 = [p.text for p in paras1]
    texts2 = [p.text for p in paras2]

    matcher = SequenceMatcher(None, texts1, texts2)
    paragraph_pairs = []

    stats = {
        'added_lines': 0,
        'removed_lines': 0,
        'modified_lines': 0,
        'unchanged_lines': 0,
    }

    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'equal':
            for idx in range(i1, i2):
                paragraph_pairs.append({
                    'type': 'equal',
                    'original': paras1[idx],
                    'comparison': paras1[idx],
                })
                stats['unchanged_lines'] += 1
        elif op == 'replace':
            orig_count = i2 - i1
            comp_count = j2 - j1
            paired = max(orig_count, comp_count)
            for k in range(paired):
                orig_p = paras1[i1 + k] if k < orig_count else None
                comp_p = paras2[j1 + k] if k < comp_count else None
                paragraph_pairs.append({
                    'type': 'modified',
                    'original': orig_p,
                    'comparison': comp_p,
                })
                stats['modified_lines'] += 1
        elif op == 'delete':
            for idx in range(i1, i2):
                paragraph_pairs.append({
                    'type': 'deleted',
                    'original': paras1[idx],
                    'comparison': None,
                })
                stats['removed_lines'] += 1
        elif op == 'insert':
            for idx in range(j1, j2):
                paragraph_pairs.append({
                    'type': 'added',
                    'original': None,
                    'comparison': paras2[idx],
                })
                stats['added_lines'] += 1

    diff_dir = os.path.join(
        os.path.dirname(docx1_path),
        "redlining",
        datetime.now().strftime("%Y%m%d_%H%M%S")
    )
    os.makedirs(diff_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(docx1_path))[0]
    comp_name = os.path.splitext(os.path.basename(docx2_path))[0]

    redline_docx_path = os.path.join(diff_dir, f"redline_{base_name}_vs_{comp_name}.docx")
    generate_redline_docx(paragraph_pairs, docx1_path, redline_docx_path, author)

    diff_html = generate_diff_html(
        paragraph_pairs,
        os.path.basename(docx1_path),
        os.path.basename(docx2_path),
        author=author,
        stats=stats,
    )

    return {
        "redline_docx_path": redline_docx_path,
        "diff_html": diff_html,
        "stats": stats,
        "paragraph_pairs": paragraph_pairs,
    }


def generate_diff_html(paragraph_pairs, file1_name, file2_name, author="Redlining System", stats=None):
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    added = stats.get('added_lines', 0) if stats else 0
    removed = stats.get('removed_lines', 0) if stats else 0
    modified = stats.get('modified_lines', 0) if stats else 0
    unchanged = stats.get('unchanged_lines', 0) if stats else 0
    total = added + removed + modified + unchanged

    html_parts = [f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: 'Times New Roman', Georgia, serif; font-size: 12pt; background: #fff; margin: 0; padding: 20px; color: #000; }}
  .report-header {{ border-bottom: 2px solid #333; padding-bottom: 15px; margin-bottom: 20px; }}
  .report-header h1 {{ font-size: 18pt; margin: 0 0 5px 0; color: #000; }}
  .report-header .meta {{ font-size: 10pt; color: #555; margin: 2px 0; }}
  .stats-bar {{ background: #f5f5f5; border: 1px solid #ddd; padding: 10px 15px; margin-bottom: 20px; font-size: 10pt; }}
  .stats-bar strong {{ margin-right: 15px; }}
  .stat-added {{ color: #0000FF; }}
  .stat-removed {{ color: #FF0000; }}
  .stat-modified {{ color: #800080; }}
  .stat-unchanged {{ color: #666; }}
  .legend {{ margin-bottom: 15px; font-size: 10pt; }}
  .legend span {{ margin-right: 20px; }}
  .legend .box {{ display: inline-block; width: 12px; height: 12px; margin-right: 4px; vertical-align: middle; }}
  .legend .box.added {{ background: #E8F0FE; border: 1px solid #0000FF; }}
  .legend .box.removed {{ background: #FEE8E8; border: 1px solid #FF0000; }}
  .legend .box.modified {{ background: #F0E8F8; border: 1px solid #800080; }}
  .legend .box.unchanged {{ background: #fff; border: 1px solid #ccc; }}
  .diff-content {{ }}
  .para {{ padding: 4px 10px; border-bottom: 1px solid #eee; line-height: 1.6; }}
  .para.added {{ background: #E8F0FE; border-left: 3px solid #0000FF; }}
  .para.removed {{ background: #FEE8E8; border-left: 3px solid #FF0000; }}
  .para.modified {{ background: #F0E8F8; border-left: 3px solid #800080; }}
  .para.unchanged {{ background: #fff; border-left: 3px solid #ccc; }}
  .para ins {{ color: #0000FF; text-decoration: underline; background: #E8F0FE; }}
  .para del {{ color: #FF0000; text-decoration: line-through; background: #FEE8E8; }}
  .para-label {{ font-size: 8pt; text-transform: uppercase; letter-spacing: 1px; margin-right: 8px; color: #999; }}
</style>
</head>
<body>
<div class="report-header">
  <h1>Document Comparison Report</h1>
  <p class="meta"><strong>Original:</strong> {file1_name}</p>
  <p class="meta"><strong>Comparison:</strong> {file2_name}</p>
  <p class="meta"><strong>Date:</strong> {now_str}</p>
  <p class="meta"><strong>Author:</strong> {author}</p>
</div>
<div class="stats-bar">
  <strong>Summary:</strong>
  <span class="stat-added">Added: {added}</span>
  <span class="stat-removed">Removed: {removed}</span>
  <span class="stat-modified">Modified: {modified}</span>
  <span class="stat-unchanged">Unchanged: {unchanged}</span>
  <span style="float:right;">Total paragraphs: {total}</span>
</div>
<div class="legend">
  <span><span class="box added"></span>Added</span>
  <span><span class="box removed"></span>Removed</span>
  <span><span class="box modified"></span>Modified</span>
  <span><span class="box unchanged"></span>Unchanged</span>
</div>
<div class="diff-content">
"""]

    for pair in paragraph_pairs:
        ptype = pair['type']
        css_class = ptype
        label = ptype.upper()

        if ptype == 'equal':
            text = pair['original'].text if pair.get('original') else ''
            html_parts.append(f'<div class="para {css_class}"><span class="para-label">{label}</span>{_html_escape(text)}</div>')
        elif ptype == 'modified':
            text1 = pair['original'].text if pair.get('original') else ''
            text2 = pair['comparison'].text if pair.get('comparison') else ''
            difflines = list(difflib.ndiff(text1.splitlines(keepends=True), text2.splitlines(keepends=True)))
            modified_html = _render_ndiff_inline(difflines)
            html_parts.append(f'<div class="para {css_class}"><span class="para-label">{label}</span>{modified_html}</div>')
        elif ptype == 'deleted':
            text = pair['original'].text if pair.get('original') else ''
            html_parts.append(f'<div class="para {css_class}"><span class="para-label">{label}</span><del>{_html_escape(text)}</del></div>')
        elif ptype == 'added':
            text = pair['comparison'].text if pair.get('comparison') else ''
            html_parts.append(f'<div class="para {css_class}"><span class="para-label">{label}</span><ins>{_html_escape(text)}</ins></div>')

    html_parts.append("</div></body></html>")
    return "\n".join(html_parts)


def _html_escape(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _render_ndiff_inline(difflines):
    result_parts = []
    for line in difflines:
        if not line:
            continue
        prefix = line[0]
        content = line[1:]
        escaped = _html_escape(content.rstrip('\n'))
        if prefix == '+':
            result_parts.append(f'<ins>{escaped}</ins>')
        elif prefix == '-':
            result_parts.append(f'<del>{escaped}</del>')
        else:
            if escaped:
                result_parts.append(f'<span>{escaped}</span>')
    combined = ''.join(result_parts)
    if not combined:
        combined = '&nbsp;'
    return combined


def generate_pdf_with_change_bars(redline_docx_path, output_pdf_path):
    if not shutil.which('libreoffice'):
        logger.warning("LibreOffice not found; skipping PDF generation")
        return None

    try:
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_pdf_path),
            redline_docx_path,
        ]
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0:
            logger.error(f"LibreOffice PDF conversion failed: {result.stderr}")
            return None

        expected_pdf = os.path.join(
            os.path.dirname(output_pdf_path),
            os.path.splitext(os.path.basename(redline_docx_path))[0] + '.pdf'
        )

        if os.path.exists(expected_pdf) and expected_pdf != output_pdf_path:
            os.rename(expected_pdf, output_pdf_path)

        if os.path.exists(output_pdf_path):
            logger.info(f"PDF with change bars saved to {output_pdf_path}")
            return output_pdf_path

        return None
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice PDF conversion timed out")
        return None
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        return None


def _generate_diff_html_from_lines(diff_lines, file1_name, file2_name):
    """Generates simple HTML showing additions (green) and deletions (red) from unified diff lines."""
    html_parts = [f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: 'Courier New', monospace; font-size: 14px; background: #fafafa; margin: 20px; }}
.diff-header {{ background: #eef; padding: 10px; border-radius: 5px; margin-bottom: 20px; }}
.diff-header h2 {{ margin: 0 0 5px 0; color: #333; }}
.diff-header p {{ margin: 0; color: #666; }}
.diff-content {{ background: #fff; border: 1px solid #ddd; border-radius: 5px; overflow-x: auto; }}
.diff-line {{ padding: 2px 10px; white-space: pre-wrap; font-family: 'Courier New', monospace; line-height: 1.5; }}
.diff-line.added {{ background: #e6ffe6; }}
.diff-line.removed {{ background: #ffe6e6; }}
.diff-line.header {{ background: #eef; color: #999; }}
.diff-line.context {{ background: #fff; }}
</style>
</head>
<body>
<div class="diff-header">
<h2>Document Comparison</h2>
<p>{file1_name} &larr; &rarr; {file2_name}</p>
</div>
<div class="diff-content">
"""]

    for line in diff_lines:
        css_class = "context"
        if line.startswith("+++") or line.startswith("---") or line.startswith("@@"):
            css_class = "header"
        elif line.startswith("+"):
            css_class = "added"
        elif line.startswith("-"):
            css_class = "removed"

        escaped = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        html_parts.append(f'<div class="diff-line {css_class}">{escaped}</div>')

    html_parts.append("</div></body></html>")
    return "\n".join(html_parts)


def compare_text_files(file1_path, file2_path):
    """Compare two text files using difflib and return structured diff data."""
    text1 = extract_text_from_file(file1_path)
    text2 = extract_text_from_file(file2_path)

    lines1 = text1.splitlines(keepends=True)
    lines2 = text2.splitlines(keepends=True)

    diff = difflib.unified_diff(
        lines1, lines2,
        fromfile=os.path.basename(file1_path),
        tofile=os.path.basename(file2_path),
        lineterm=""
    )

    diff_lines = list(diff)
    file1_name = os.path.basename(file1_path)
    file2_name = os.path.basename(file2_path)

    diff_html = _generate_diff_html_from_lines(diff_lines, file1_name, file2_name)

    added_lines = sum(1 for l in diff_lines if l.startswith("+") and not l.startswith("+++"))
    removed_lines = sum(1 for l in diff_lines if l.startswith("-") and not l.startswith("---"))
    unchanged_lines = sum(1 for l in diff_lines if l.startswith(" "))

    return {
        "diff_html": diff_html,
        "added_lines": added_lines,
        "removed_lines": removed_lines,
        "unchanged_lines": unchanged_lines,
    }
