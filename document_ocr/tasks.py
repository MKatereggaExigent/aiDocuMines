import os
import logging
import pandas as pd
import fitz
from celery import shared_task, chord
from django.db import transaction
from django.shortcuts import get_object_or_404
from django.utils.timezone import now
from core.models import File
import shutil
from docx import Document
# from document_ocr.models import OCRFile
from document_ocr.models import OCRRun, OCRFile
from document_ocr.utils import OCRService, cleanup_tmp_dir
import uuid
from core.utils import register_generated_file


logger = logging.getLogger(__name__)



@shared_task
def process_ocr(run_id, file_id, ocr_option):
    """
    Celery task to perform OCR on a document asynchronously.
    """
    try:
        try:
            run = OCRRun.objects.get(id=run_id)
        except OCRRun.DoesNotExist:
            logger.error(f"❌ OCRRun with id {run_id} does not exist. Skipping OCR.")
            return {"error": f"OCRRun with id {run_id} not found"}

        # ✅ Retrieve the original file using file_id (INTEGER)
        file_obj = get_object_or_404(File, id=file_id)

        # ✅ Create the OCRFile entry before processing
        ocr_file, created = OCRFile.objects.get_or_create(
            original_file=file_obj, defaults={"status": "Processing", "run": run}
        )

        run.status = "Processing"
        run.save()

        ocr_service = OCRService()

        # ✅ Skip non-PDF files
        if not ocr_service.is_pdf(file_obj.filepath):
            logger.info(f"🔹 Skipping OCR: {file_obj.filepath} is not a PDF.")
            ocr_file.status = "Completed"
            ocr_file.save()
            return {"message": "File is not a PDF. OCR skipped."}

        # ✅ Extract bookmarks before OCR processing
        bookmarks_df = ocr_service.extract_bookmarks_to_dataframe(file_obj.filepath)
        bookmarks_list = bookmarks_df.to_dict(orient="records")

        # ✅ Burst the PDF into page batches
        batch_files = ocr_service.burst_pdf(file_obj.filepath)

        if not batch_files:
            ocr_file.status = "Failed"
            ocr_file.save()
            return {"error": "Failed to burst the PDF", "file_path": file_obj.filepath}

        # ✅ Define OCR tasks for each batch
        ocr_tasks = [
            ocr_pdf_page_batch.s(ocr_file.id, batch_file, start_page, end_page, ocr_option)  # 🔥 Use `ocr_file.id` (UUID)
            for start_page, end_page, batch_file in batch_files
        ]

        # ✅ Wait for all OCR tasks before merging
        workflow = chord(ocr_tasks)(merge_ocr_batches.s(ocr_file.id, bookmarks_list))  # 🔥 Use `ocr_file.id` (UUID)

        return workflow

    except Exception as e:
        logger.error(f"❌ OCR processing failed: {e}")
        if 'ocr_file' in locals():
            ocr_file.status = "Failed"
            ocr_file.save()
        return {"error": str(e)}


@shared_task
def ocr_pdf_file(file_id, ocr_option="basic"):
    """Main OCR task that processes a PDF file and applies OCR."""
    file_entry = get_object_or_404(OCRFile, id=file_id)

    # ✅ Instantiate OCR Service
    ocr_service = OCRService()

    # ✅ Ensure the file exists and is a PDF
    file_path = ocr_service.get_file_path(file_entry.original_file.id)
    if not file_path or not ocr_service.is_pdf(file_path):
        logger.info(f"🔹 Skipping OCR: File {file_path} is not a PDF or does not exist.")
        return None

    # ✅ Get the OCR directory
    ocr_dir = os.path.join(os.path.dirname(file_entry.original_file.filepath), "ocr")

    # ✅ Ensure OCR folder exists
    if not os.path.exists(ocr_dir):
        os.makedirs(ocr_dir, exist_ok=True)

    # ✅ Check if an OCR file already exists before proceeding
    existing_ocr_files = [f for f in os.listdir(ocr_dir) if f.endswith(".pdf") and "ocr-" in f]

    if existing_ocr_files:
        existing_ocr_file = os.path.join(ocr_dir, existing_ocr_files[0])
        logger.info(f"⚠️ Skipping OCR: Existing OCR file found at {existing_ocr_file}")

        with transaction.atomic():
            file_entry.ocr_filepath = existing_ocr_file
            file_entry.status = "Completed"
            file_entry.updated_at = now()
            file_entry.save()

        return {"message": "OCR already completed", "ocr_file": existing_ocr_file}

    # ✅ Proceed with OCR if no existing file is found
    logger.info(f"🔄 Starting OCR for {file_path}")

    # ✅ Step 1: Extract bookmarks before processing and store them
    bookmarks_df = ocr_service.extract_bookmarks_to_dataframe(file_path)
    bookmarks_list = bookmarks_df.to_dict(orient='records')  # Convert DataFrame to list of dicts for Celery serialization

    # ✅ Step 2: Burst the PDF into smaller page batches
    batch_files = ocr_service.burst_pdf(file_path)

    if not batch_files:
        file_entry.status = "Failed"
        file_entry.save()
        return {"error": "Failed to burst the PDF", "file_path": file_path}

    # ✅ Step 3: Define OCR tasks for each batch
    ocr_tasks = [
        ocr_pdf_page_batch.s(file_id, batch_file, start_page, end_page, ocr_option)
        for start_page, end_page, batch_file in batch_files
    ]

    # ✅ Step 4: Use `chord` to wait for OCR tasks before merging
    callback = merge_ocr_batches.s(file_id, bookmarks_list)  # 🔹 Ensure bookmarks_list is passed here

    workflow = chord(ocr_tasks)(callback)
    return workflow


@shared_task
def ocr_pdf_page_batch(file_id, batch_file_path, start_page, end_page, ocr_option="basic"):
    """OCR task for a batch of pages in a PDF."""
    try:
        # ✅ Ensure the file exists before proceeding
        batch_file_path = os.path.normpath(batch_file_path)  # ✅ Fix double `/ocr/tmp/` issue
        if not os.path.exists(batch_file_path):
            logger.error(f"❌ Batch file not found: {batch_file_path}")
            return {"error": "Batch file not found", "batch_file_path": batch_file_path}

        # ✅ Instantiate OCR Service
        ocr_service = OCRService()

        # Apply OCR to the batch file
        ocr_file = ocr_service.apply_ocr(file_id, batch_file_path, ocr_option)
        return {"start_page": start_page, "end_page": end_page, "ocr_file": ocr_file}

    except Exception as e:
        logger.error(f"❌ OCR processing failed: {str(e)}")
        return {"error": str(e), "batch_file_path": batch_file_path}
    
    
@shared_task
def merge_ocr_batches(results, ocr_file_id, bookmarks_list):
    """Merges all OCR'ed PDF batches and reattaches bookmarks."""
    file_entry = get_object_or_404(OCRFile, id=ocr_file_id)

    # ✅ Ensure all paths are correct
    ocr_dir = os.path.join(os.path.dirname(file_entry.original_file.filepath), "ocr")
    tmp_dir = os.path.join(ocr_dir, "tmp")

    ocr_service = OCRService()

    # ✅ Filter out `None` OCR outputs
    ocr_files = [res["ocr_file"] for res in results if res.get("ocr_file") is not None]

    if not ocr_files:
        logger.error(f"❌ No valid OCR'ed pages found. Merging failed.")
        file_entry.status = "Failed"
        file_entry.save()
        return None

    # ✅ Merge final OCR document inside `ocr/`
    final_pdf_path = os.path.join(ocr_dir, f"ocr-{uuid.uuid4()}.pdf")
    merged_pdf = ocr_service.merge_pdf(ocr_files, final_pdf_path)

    if not merged_pdf:
        logger.error(f"❌ Merging failed. No final PDF created.")
        file_entry.status = "Failed"
        file_entry.save()
        return None


        # Register OCR’d PDF in File table
        upload_run = file_entry.original_file.run
        registered = register_generated_file(
            file_path=final_pdf_path,
            user=file_entry.original_file.user,
            run=upload_run,
            project_id=file_entry.original_file.project_id,
            service_id=file_entry.original_file.service_id,
            folder_name="ocr"
        )


    # ✅ Reattach bookmarks
    pdf_document = fitz.open(final_pdf_path)
    total_pages = pdf_document.page_count
    pdf_document.close()

    bookmarks_df = pd.DataFrame(bookmarks_list)

    try:
        ocr_service.reattach_bookmarks_from_dataframe(final_pdf_path, bookmarks_df, 1, total_pages)
    except Exception as e:
        logger.error(f"❌ Failed to reattach bookmarks: {e}")

    # ✅ Update database record
    with transaction.atomic():
        file_entry.ocr_filepath = final_pdf_path
        file_entry.status = "Processed"
        file_entry.updated_at = now()
        file_entry.save()

    # ✅ Delete `tmp/` after merging
    shutil.rmtree(tmp_dir, ignore_errors=True)
    logger.info(f"✅ Cleaned up temporary directory: {tmp_dir}")

    # ✅ **Trigger DOCX conversion**
    process_pdf_to_docx.delay(file_entry.id, final_pdf_path)

    # return final_pdf_path


    return {
        "ocr_run_id": str(file_entry.run.id),
        "file_id": file_entry.original_file.id,
        "ocr_file_id": file_entry.id,
        "ocr_merged_pdf": final_pdf_path,
        "status": "Completed",
        "registered_outputs": [{
            "file_id": registered.id,
            "filename": registered.filename,
            "path": registered.filepath
        }]
    }


@shared_task
def process_pdf_to_docx(ocr_file_id, final_pdf_path):
    """Converts an OCR’ed PDF to formatted DOCX and registers it in the File model."""
    file_entry = get_object_or_404(OCRFile, id=ocr_file_id)

    # ✅ Ensure OCRed PDF exists before conversion
    if not final_pdf_path or not os.path.exists(final_pdf_path):
        logger.error(f"❌ No OCRed PDF found at {final_pdf_path}. Skipping DOCX conversion.")
        return {"error": "OCRed PDF not found", "ocr_file_id": ocr_file_id}

    ocr_service = OCRService()
    ocr_dir = os.path.dirname(final_pdf_path)
    formatted_docx_path = os.path.join(ocr_dir, f"{uuid.uuid4()}_formatted.docx")

    # ✅ Convert to formatted DOCX
    formatted_output = ocr_service.convert_to_formatted_docx(final_pdf_path, formatted_docx_path)

    registered = None
    if formatted_output:
        with transaction.atomic():
            file_entry.docx_path = formatted_output
            file_entry.save()

        # ✅ Register in File table
        registered = register_generated_file(
            file_path=formatted_output,
            user=file_entry.original_file.user,
            run=file_entry.original_file.run,
            project_id=file_entry.original_file.project_id,
            service_id=file_entry.original_file.service_id,
            folder_name="ocr"
        )

        # ✅ Trigger raw DOCX generation
        generate_raw_docx.delay(file_entry.id)

    return {
        "ocr_run_id": str(file_entry.run.id),
        "file_id": file_entry.original_file.id,
        "ocr_file_id": file_entry.id,
        "formatted_docx": formatted_output,
        "status": "Completed" if formatted_output else "Failed",
        "registered_outputs": [{
            "file_id": registered.id,
            "filename": registered.filename,
            "path": registered.filepath
        }] if registered else []
    }



'''
@shared_task
def process_pdf_to_docx(ocr_file_id, final_pdf_path):
    """Converts an OCR’ed PDF to both formatted and raw DOCX and updates the database."""
    file_entry = get_object_or_404(OCRFile, id=ocr_file_id)

    # ✅ Ensure OCRed PDF exists before conversion
    if not final_pdf_path or not os.path.exists(final_pdf_path):
        logger.error(f"❌ No OCRed PDF found at {final_pdf_path}. Skipping DOCX conversion.")
        return None

    ocr_service = OCRService()
    
    # ✅ **Define output paths inside `ocr/` directory**
    ocr_dir = os.path.dirname(final_pdf_path)
    formatted_docx_path = os.path.join(ocr_dir, f"{uuid.uuid4()}_formatted.docx")

    # ✅ **Convert to formatted DOCX using Adobe**
    formatted_output = ocr_service.convert_to_formatted_docx(final_pdf_path, formatted_docx_path)

    # ✅ **Update database with formatted DOCX path**
    with transaction.atomic():
        if formatted_output:
            file_entry.docx_path = formatted_output
            file_entry.save()

    # ✅ **Trigger raw DOCX generation**
    if formatted_output:
        generate_raw_docx.delay(file_entry.id)

    register_generated_file(
        file_path=formatted_output,
        user=file_entry.original_file.user,
        run=file_entry.original_file.run,
        project_id=file_entry.original_file.project_id,
        service_id=file_entry.original_file.service_id,
        folder_name="ocr"
    )


    return {"formatted_docx": formatted_output}
'''


@shared_task
def generate_raw_docx(ocr_file_id):
    """Generate and register a raw DOCX with unformatted text extracted from the formatted DOCX."""
    file_entry = get_object_or_404(OCRFile, id=ocr_file_id)

    if not file_entry.docx_path or not os.path.exists(file_entry.docx_path):
        logger.error(f"❌ No formatted DOCX found for ocr_file_id: {ocr_file_id}. Skipping RAW DOCX creation.")
        return {
            "ocr_file_id": ocr_file_id,
            "status": "Failed",
            "error": "Formatted DOCX not found"
        }

    # ✅ Define path for RAW DOCX
    ocr_dir = os.path.dirname(file_entry.docx_path)
    raw_docx_path = os.path.join(ocr_dir, f"{uuid.uuid4()}_raw.docx")

    original_doc = Document(file_entry.docx_path)
    raw_doc = Document()

    for para in original_doc.paragraphs:
        if para.text.strip():
            new_para = raw_doc.add_paragraph(para.text)

            new_para.paragraph_format.left_indent = None
            new_para.paragraph_format.right_indent = None
            new_para.paragraph_format.first_line_indent = None
            new_para.paragraph_format.space_before = None
            new_para.paragraph_format.space_after = None
            new_para.paragraph_format.alignment = None

            for run in new_para.runs:
                run.bold = False
                run.italic = False
                run.underline = False

    raw_doc.save(raw_docx_path)

    # ✅ Register file
    registered = register_generated_file(
        file_path=raw_docx_path,
        user=file_entry.original_file.user,
        run=file_entry.original_file.run,
        project_id=file_entry.original_file.project_id,
        service_id=file_entry.original_file.service_id,
        folder_name="ocr"
    )

    # ✅ Save to DB
    with transaction.atomic():
        file_entry.raw_docx_path = raw_docx_path
        file_entry.save()

    logger.info(f"✅ Raw DOCX created and registered: {raw_docx_path}")

    return {
        "ocr_run_id": str(file_entry.run.id),
        "file_id": file_entry.original_file.id,
        "ocr_file_id": file_entry.id,
        "raw_docx": raw_docx_path,
        "status": "Completed",
        "registered_outputs": [{
            "file_id": registered.id,
            "filename": registered.filename,
            "path": registered.filepath
        }]
    }




'''
@shared_task
def generate_raw_docx(ocr_file_id):
    """Generate a raw DOCX file with unformatted text extracted from the formatted DOCX."""
    file_entry = get_object_or_404(OCRFile, id=ocr_file_id)

    # ✅ Ensure that a formatted DOCX exists before proceeding
    if not file_entry.docx_path or not os.path.exists(file_entry.docx_path):
        logger.error(f"❌ No formatted DOCX found for ocr_file_id: {ocr_file_id}. Skipping RAW DOCX creation.")
        return None

    # ✅ Define path for the RAW DOCX in the same `ocr/` directory
    ocr_dir = os.path.dirname(file_entry.docx_path)
    raw_docx_path = os.path.join(ocr_dir, f"{uuid.uuid4()}_raw.docx")

    # ✅ Load the formatted DOCX
    original_doc = Document(file_entry.docx_path)
    raw_doc = Document()  # Create a new blank DOCX

    # ✅ Extract text without any formatting
    for para in original_doc.paragraphs:
        if para.text.strip():  # Check if the paragraph has non-blank text
            new_para = raw_doc.add_paragraph(para.text)

            # ✅ Remove any potential numbering or indentation by resetting paragraph formatting
            new_para.paragraph_format.left_indent = None
            new_para.paragraph_format.right_indent = None
            new_para.paragraph_format.first_line_indent = None
            new_para.paragraph_format.space_before = None
            new_para.paragraph_format.space_after = None
            new_para.paragraph_format.alignment = None

            # ✅ Ensure no text is bolded, italicized, or underlined
            for run in new_para.runs:
                run.bold = False
                run.italic = False
                run.underline = False

    # ✅ Save the raw DOCX
    raw_doc.save(raw_docx_path)

    register_generated_file(
        file_path=raw_docx_path,
        user=file_entry.original_file.user,
        run=file_entry.original_file.run,
        project_id=file_entry.original_file.project_id,
        service_id=file_entry.original_file.service_id,
        folder_name="ocr"
    )

    # ✅ Update database record with raw_docx path
    with transaction.atomic():
        file_entry.raw_docx_path = raw_docx_path
        file_entry.save()

    logger.info(f"✅ Raw DOCX created at {raw_docx_path}")

    return raw_docx_path

'''
